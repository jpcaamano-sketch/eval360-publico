"""
Evaluación 360 v2 — App Admin
Puerto: 8506
"""

import io
import streamlit as st
import streamlit.components.v1 as components
from datetime import datetime, timezone
from core.config import ESCALA, GOOGLE_API_KEY
from core.styles import ADMIN_CSS
from core import queries
from core.email_service import (
    enviar_invitacion_autoevaluacion,
    enviar_invitacion_feedback,
    enviar_recordatorio_autoevaluacion,
    enviar_recordatorio_feedback,
)
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

st.set_page_config(page_title="Evaluación 360° Admin", page_icon="📊", layout="wide")
st.markdown(ADMIN_CSS, unsafe_allow_html=True)


# ============================================================
# SIDEBAR
# ============================================================

MENU_OPTIONS = [
    "Inicio",
    "Personas y Empresas",
    "Ingreso Encuestas",
    "Ingreso de Grupos",
    "Seguimiento Autoevaluaciones",
    "Seguimiento Feedback",
    "Informe Final 360",
    "Ingresos Especiales",
]

with st.sidebar:
    st.markdown('<a href="http://localhost:8520" target="_self" style="display:block;color:rgba(255,255,255,0.75);font-size:0.82rem;text-decoration:none;margin-bottom:6px;">← Volver al Hub</a>', unsafe_allow_html=True)
    st.title("Evaluación 360°")
    st.divider()

    menu = st.radio(
        "nav",
        MENU_OPTIONS,
        label_visibility="collapsed",
        key="nav_radio",
    )

    st.divider()

# JS: colorea el último ítem del radio (Ingresar Evaluadores) en ámbar
components.html("""
<script>
function styleLastOption() {
    try {
        var doc = window.parent.document;
        var labels = doc.querySelectorAll(
            '[data-testid="stSidebar"] [data-testid="stRadio"] label'
        );
        if (labels.length === 0) return;
        var last = labels[labels.length - 1];
        last.style.setProperty('color', '#FFA040', 'important');
        last.style.setProperty('font-style', 'italic', 'important');
        last.querySelectorAll('p, span, div').forEach(function(el) {
            el.style.setProperty('color', '#FFA040', 'important');
            el.style.setProperty('font-style', 'italic', 'important');
        });
    } catch(e) {}
}
styleLastOption();
setTimeout(styleLastOption, 300);
setTimeout(styleLastOption, 800);
new MutationObserver(styleLastOption).observe(
    window.parent.document.body, {childList: true, subtree: true}
);
</script>
""", height=0)


# ============================================================
# PÁGINA: INICIO
# ============================================================

def pagina_inicio():
    st.header("Inicio")
    st.markdown("Bienvenido al panel de administración de **Evaluación 360° v2**.")

    col1, col2, col3 = st.columns(3)

    plantillas = queries.listar_plantillas()
    grupos = queries.listar_grupos()

    with col1:
        st.metric("Plantillas", len(plantillas))
    with col2:
        st.metric("Grupos", len(grupos))
    with col3:
        total_part = 0
        for g in grupos:
            parts = queries.listar_participantes(g["id"])
            total_part += len(parts)
        st.metric("Participantes", total_part)

    st.divider()
    st.subheader("Flujo del proceso")
    st.markdown("""
    1. **Crear Plantilla** con categorías y competencias
    2. **Crear Grupo** y asignarle una plantilla
    3. **Agregar Participantes** al grupo
    4. **Enviar Invitación de Autoevaluación** a los participantes
    5. Los participantes completan su autoevaluación e ingresan sus evaluadores
    6. **Enviar Invitación de Feedback** a los evaluadores
    7. **Seguimiento** del progreso de autoevaluaciones y feedback
    """)


# ============================================================
# PÁGINA: PLANTILLAS
# ============================================================

def _vista_previa_plantilla(plantilla_id, categorias):
    """Muestra resumen visual de categorías y competencias."""
    with st.container(border=True):
        if not categorias:
            st.caption("Sin categorías ni competencias.")
            return
        for cat in categorias:
            comps = queries.listar_competencias(cat["id"])
            st.markdown(f"**📁 {cat['nombre']}** ({len(comps)})")
            if comps:
                for comp in comps:
                    st.markdown(
                        f"&nbsp;&nbsp;&nbsp;&nbsp;• _{comp['texto_auto']}_ &nbsp;/&nbsp; {comp['texto_feedback']}"
                    )
            else:
                st.caption("    Sin competencias")


def _editar_plantilla(plantilla_id):
    """Vista tipo planilla para editar una plantilla."""
    plantilla = queries.obtener_plantilla(plantilla_id)
    if not plantilla:
        st.error("Plantilla no encontrada")
        st.session_state["plantilla_editando"] = None
        return

    # Cabecera
    col_back, col_title = st.columns([1, 5])
    with col_back:
        if st.button("← Volver"):
            st.session_state["plantilla_editando"] = None
            st.session_state["editando_comp"] = None
            st.rerun()
    with col_title:
        st.subheader(plantilla["nombre"])

    # Editar nombre/desc inline
    with st.expander("Editar nombre y descripción"):
        with st.form("edit_plantilla_datos"):
            c1, c2 = st.columns(2)
            with c1:
                nuevo_nombre = st.text_input("Nombre", value=plantilla["nombre"])
            with c2:
                nueva_desc = st.text_input("Descripción", value=plantilla.get("descripcion") or "")
            if st.form_submit_button("Guardar"):
                queries.actualizar_plantilla(plantilla_id, {
                    "nombre": nuevo_nombre.strip(),
                    "descripcion": nueva_desc.strip(),
                })
                st.rerun()

    st.divider()

    # --- Categorías: agregar inline ---
    categorias = queries.listar_categorias(plantilla_id)
    cat_nombres = [c["nombre"] for c in categorias]

    st.markdown("**Categorías**")
    # Mostrar categorías como chips con delete
    if categorias:
        cols_cat = st.columns(min(len(categorias) + 1, 6))
        for i, cat in enumerate(categorias):
            with cols_cat[i % (min(len(categorias) + 1, 6))]:
                st.markdown(f"📁 **{cat['nombre']}**")
                if st.button("✕", key=f"delcat_{cat['id']}", help="Eliminar categoría"):
                    queries.eliminar_categoria(cat["id"])
                    st.rerun()

    with st.form("nueva_cat"):
        c1, c2 = st.columns([3, 1])
        with c1:
            nueva_cat = st.text_input("Nueva categoría", placeholder="Nombre de categoría")
        with c2:
            st.markdown("")
            agregar_cat = st.form_submit_button("➕ Agregar")
        if agregar_cat and nueva_cat.strip():
            queries.crear_categoria(plantilla_id, nueva_cat.strip(), len(categorias) + 1)
            st.rerun()

    st.divider()

    # --- Tabla de competencias ---
    st.markdown("**Competencias**")

    # Cabecera de la tabla
    h1, h2, h3, h4, h5 = st.columns([1.5, 2.5, 2.5, 0.5, 0.5])
    with h1:
        st.markdown("**Categoría**")
    with h2:
        st.markdown("**Texto 1ª persona (auto)**")
    with h3:
        st.markdown("**Texto 3ª persona (feedback)**")
    with h4:
        st.markdown("**✏️**")
    with h5:
        st.markdown("**🗑️**")

    st.markdown("---")

    # Filas de competencias
    todas = queries.listar_competencias_por_plantilla(plantilla_id)

    if not todas and not categorias:
        st.info("Agrega categorías y luego competencias.")
    elif not todas:
        st.info("No hay competencias. Agrega una abajo.")

    for comp in todas:
        # Si estamos editando esta competencia, mostrar inputs
        if st.session_state.get("editando_comp") == comp["id"]:
            with st.form(f"edit_{comp['id']}"):
                e1, e2, e3, e4, e5 = st.columns([1.5, 2.5, 2.5, 0.5, 0.5])
                with e1:
                    # Selector de categoría
                    idx_actual = cat_nombres.index(comp["categoria_nombre"]) if comp["categoria_nombre"] in cat_nombres else 0
                    nueva_cat_comp = st.selectbox("Cat", options=categorias, format_func=lambda c: c["nombre"],
                                                  index=idx_actual, label_visibility="collapsed", key=f"ecat_{comp['id']}")
                with e2:
                    edit_auto = st.text_area("Auto", value=comp["texto_auto"], label_visibility="collapsed",
                                             key=f"eauto_{comp['id']}", height=80)
                with e3:
                    edit_feed = st.text_area("Feed", value=comp["texto_feedback"], label_visibility="collapsed",
                                             key=f"efeed_{comp['id']}", height=80)
                with e4:
                    guardar = st.form_submit_button("💾")
                with e5:
                    pass
                if guardar:
                    queries.actualizar_competencia(comp["id"], {
                        "texto_auto": edit_auto.strip(),
                        "texto_feedback": edit_feed.strip(),
                        "categoria_id": nueva_cat_comp["id"],
                    })
                    st.session_state["editando_comp"] = None
                    st.rerun()
            # Botón cancelar fuera del form
            if st.button("Cancelar", key=f"cancel_{comp['id']}"):
                st.session_state["editando_comp"] = None
                st.rerun()
        else:
            r1, r2, r3, r4, r5 = st.columns([1.5, 2.5, 2.5, 0.5, 0.5])
            with r1:
                st.caption(comp["categoria_nombre"])
            with r2:
                st.markdown(comp["texto_auto"])
            with r3:
                st.markdown(comp["texto_feedback"])
            with r4:
                if st.button("✏️", key=f"editc_{comp['id']}"):
                    st.session_state["editando_comp"] = comp["id"]
                    st.rerun()
            with r5:
                if st.button("🗑️", key=f"delc_{comp['id']}"):
                    queries.eliminar_competencia(comp["id"])
                    st.rerun()

    # --- Agregar nueva competencia ---
    st.markdown("---")
    st.markdown("**Agregar competencia**")
    if not categorias:
        st.warning("Primero crea al menos una categoría.")
    else:
        with st.form("nueva_comp"):
            n1, n2, n3, n4 = st.columns([1.5, 2.5, 2.5, 1])
            with n1:
                cat_sel = st.selectbox("Categoría", options=categorias,
                                       format_func=lambda c: c["nombre"], label_visibility="collapsed")
            with n2:
                txt_auto = st.text_area("Texto 1ª persona", placeholder="Ej: Comunico mis ideas con claridad",
                                        label_visibility="collapsed", height=80)
            with n3:
                txt_feed = st.text_area("Texto 3ª persona", placeholder="Ej: Comunica sus ideas con claridad",
                                        label_visibility="collapsed", height=80)
            with n4:
                st.markdown("")
                agregar = st.form_submit_button("➕ Agregar")
            if agregar:
                if txt_auto.strip() and txt_feed.strip():
                    comps_cat = queries.listar_competencias(cat_sel["id"])
                    queries.crear_competencia(cat_sel["id"], txt_auto.strip(), txt_feed.strip(), len(comps_cat) + 1)
                    st.rerun()
                else:
                    st.warning("Completa ambos textos")


# ============================================================
# PÁGINA: GRUPOS
# ============================================================

def pagina_grupos():
    if "grupo_detalle" not in st.session_state:
        st.session_state["grupo_detalle"] = None

    st.header("Ingreso de Grupos")

    tab_lista, tab_nuevo = st.tabs(["Grupos existentes", "Crear nuevo grupo"])

    with tab_nuevo:
        import pandas as pd
        plantillas_n = sorted(queries.listar_plantillas(solo_activas=True), key=lambda p: (p["nombre"] or "").lower())
        empresas_otec_n = queries.listar_empresas_otec()

        if not plantillas_n:
            st.warning("Primero crea una plantilla antes de crear un grupo.")
        elif not empresas_otec_n:
            st.warning("No hay empresas registradas. Crea una en Gestión OTEC primero.")
        else:
            col1, col2, col3 = st.columns(3)
            with col1:
                empresa_sel_n = st.selectbox(
                    "Empresa",
                    options=[None] + empresas_otec_n,
                    format_func=lambda e: e["nombre_empresa"] if e else "— Selecciona empresa —",
                    key="ng_empresa",
                )
            with col2:
                nombre_n = st.text_input("Nombre del grupo", key="ng_nombre")
            with col3:
                plantilla_sel_n = st.selectbox(
                    "Plantilla",
                    options=[None] + plantillas_n,
                    format_func=lambda p: p["nombre"] if p else "— Selecciona plantilla —",
                    key="ng_plantilla",
                )

            rut_empresa_n = empresa_sel_n["rut_empresa"] if empresa_sel_n else None
            seleccionados_n = []
            personas_n = []
            rows_n = []

            # Mostrar tabla solo cuando empresa Y plantilla estén seleccionadas
            if empresa_sel_n and plantilla_sel_n:
                personas_n = queries.listar_personas_sist(rut_empresa=rut_empresa_n)
                if not personas_n:
                    personas_n = queries.listar_personas_sist()

                st.divider()
                st.markdown("**Selecciona los participantes:**")

                if not personas_n:
                    st.info("No hay personas con correo registradas en sist_personas.")
                else:
                    personas_n_ord = sorted(personas_n, key=lambda p: (p.get("pers_apellidos") or "").lower())
                    rows_n = [
                        {
                            "✓": False,
                            "#": i,
                            "RUT": p["pers_rut"],
                            "Nombres": p.get("pers_nombres", ""),
                            "Apellidos": p.get("pers_apellidos", ""),
                            "Correo": p.get("pers_correo", ""),
                        }
                        for i, p in enumerate(personas_n_ord, 1)
                    ]
                    df_n = pd.DataFrame(rows_n)

                    edited_n = st.data_editor(
                        df_n,
                        column_config={
                            "✓": st.column_config.CheckboxColumn("✓", default=False, width="small"),
                            "#": st.column_config.NumberColumn("#", disabled=True, width="small"),
                            "RUT": st.column_config.TextColumn("RUT", disabled=True, width=120),
                            "Nombres": st.column_config.TextColumn("Nombres", disabled=True, width=160),
                            "Apellidos": st.column_config.TextColumn("Apellidos", disabled=True, width=160),
                            "Correo": st.column_config.TextColumn("Correo", disabled=True),
                        },
                        disabled=["#", "RUT", "Nombres", "Apellidos", "Correo"],
                        hide_index=True,
                        use_container_width=True,
                        key=f"ng_editor_{rut_empresa_n or 'todos'}",
                    )

                    seleccionados_n = edited_n.index[edited_n["✓"] == True].tolist()
                    st.caption(f"{len(seleccionados_n)} persona(s) seleccionada(s) de {len(personas_n_ord)}")
            else:
                st.info("Selecciona empresa y plantilla para ver los participantes disponibles.")

            st.divider()
            if st.button("💾 Guardar selección", use_container_width=True, type="primary", key="ng_guardar"):
                errores_n = []
                if not empresa_sel_n:
                    errores_n.append("Selecciona una empresa.")
                if not plantilla_sel_n:
                    errores_n.append("Selecciona una plantilla.")
                if not nombre_n.strip():
                    errores_n.append("Ingresa un nombre de grupo.")
                if personas_n and len(seleccionados_n) == 0:
                    errores_n.append("Selecciona al menos un participante.")
                if errores_n:
                    for e in errores_n:
                        st.warning(e)
                else:
                    try:
                        grupo_nuevo = queries.crear_grupo(nombre_n.strip(), plantilla_sel_n["id"], rut_empresa_n)
                        gid = grupo_nuevo["id"]
                        for idx in seleccionados_n:
                            queries.crear_participante(gid, rows_n[idx]["RUT"])
                        st.success(f"✅ Grupo '{nombre_n.strip()}' creado con {len(seleccionados_n)} participante(s).")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")

    with tab_lista:
        import pandas as pd
        grupos = queries.listar_grupos()
        if not grupos:
            st.info("No hay grupos. Crea uno en la pestaña 'Crear nuevo grupo'.")
            return

        rows_orig = []
        for i, g in enumerate(grupos, 1):
            plantilla_nombre = g.get("v2_plantillas", {}).get("nombre", "—") if g.get("v2_plantillas") else "—"
            n_parts = len(queries.listar_participantes(g["id"]))
            rows_orig.append({
                "✓": False,
                "#": i,
                "Empresa": g.get("empresa") or "",
                "Nombre": g.get("nombre") or "",
                "Plantilla": plantilla_nombre,
                "Estado": g.get("estado") or "",
                "Participantes": n_parts,
            })

        st.caption("Marca ✓ un grupo y pulsa **Editar** para ver y agregar participantes.")
        df_orig = pd.DataFrame(rows_orig)

        edited_df = st.data_editor(
            df_orig,
            column_config={
                "✓": st.column_config.CheckboxColumn("✓", default=False, width="small"),
                "#": st.column_config.NumberColumn("#", disabled=True, width="small"),
                "Empresa": st.column_config.TextColumn("Empresa", disabled=True),
                "Nombre": st.column_config.TextColumn("Nombre"),
                "Plantilla": st.column_config.TextColumn("Plantilla", disabled=True),
                "Estado": st.column_config.TextColumn("Estado", disabled=True),
                "Participantes": st.column_config.NumberColumn("Participantes", disabled=True, width="small"),
            },
            disabled=["#", "Empresa", "Plantilla", "Estado", "Participantes"],
            hide_index=True,
            use_container_width=True,
            key="grupos_editor",
        )

        a_sel = edited_df.index[edited_df["✓"] == True].tolist()
        col_abrir, col_grabar, col_eliminar, _ = st.columns([1.2, 1.2, 1.2, 2.8])

        with col_abrir:
            if st.button(
                "Editar",
                use_container_width=True,
                type="primary",
                disabled=(len(a_sel) != 1),
                help="Marca un grupo para editarlo y agregar participantes",
            ):
                st.session_state["grupo_detalle"] = grupos[a_sel[0]]["id"]
                st.rerun()

        with col_grabar:
            if st.button("Grabar", use_container_width=True):
                cambios = 0
                for idx, (orig, edit) in enumerate(zip(rows_orig, edited_df.to_dict("records"))):
                    if edit["Nombre"] != orig["Nombre"]:
                        queries.actualizar_grupo(grupos[idx]["id"], {
                            "nombre": edit["Nombre"].strip(),
                        })
                        cambios += 1
                if cambios:
                    st.success(f"✅ {cambios} grupo(s) actualizado(s).")
                    st.rerun()
                else:
                    st.info("No hay cambios para grabar.")

        with col_eliminar:
            if st.button(
                f"Eliminar ({len(a_sel)})" if a_sel else "Eliminar",
                use_container_width=True,
                disabled=(len(a_sel) == 0),
            ):
                for idx in a_sel:
                    queries.eliminar_grupo(grupos[idx]["id"])
                st.success(f"✅ {len(a_sel)} grupo(s) eliminado(s).")
                st.rerun()

        if st.session_state["grupo_detalle"]:
            st.divider()
            col_det, col_cerrar = st.columns([5, 1])
            with col_cerrar:
                if st.button("Cerrar", use_container_width=True):
                    st.session_state["grupo_detalle"] = None
                    st.rerun()
            _detalle_grupo(st.session_state["grupo_detalle"])



def _detalle_grupo(grupo_id):
    """Vista de detalle de un grupo."""
    grupo = queries.obtener_grupo(grupo_id)
    if not grupo:
        st.error("Grupo no encontrado")
        st.session_state["grupo_detalle"] = None
        return

    import pandas as pd

    participantes = queries.listar_participantes(grupo_id)
    huerfanos = [p for p in participantes if not p.get("pers_rut")]
    rut_empresa_grupo = grupo.get("rut_empresa")

    if huerfanos:
        st.warning(
            f"⚠️ **{len(huerfanos)} participante(s) sin vincular** a sist_personas (datos previos a la migración). "
            "Elimínalos o vincúlalos desde Personas y Empresas."
        )

    # Cargar personas disponibles (filtradas por empresa, fallback todas)
    todas_personas = queries.listar_personas_sist(rut_empresa=rut_empresa_grupo)
    if not todas_personas:
        todas_personas = queries.listar_personas_sist()

    # Conteo de evaluadores por participante (una sola consulta)
    ev_list = queries.listar_evaluadores_por_grupo(grupo_id)
    ev_count = {}
    for ev in ev_list:
        pid = ev.get("participante_id_ref")
        if pid:
            ev_count[pid] = ev_count.get(pid, 0) + 1

    enrolled_ruts = {p["pers_rut"] for p in participantes if p.get("pers_rut")}
    part_by_rut = {p["pers_rut"]: p for p in participantes if p.get("pers_rut")}

    # Construir filas: todas las personas ordenadas por apellido, pre-marcar las inscritas
    todas_personas_ord = sorted(todas_personas, key=lambda p: (p.get("pers_apellidos") or "").lower())
    rows_part = []
    for i, per in enumerate(todas_personas_ord, 1):
        rut = per["pers_rut"]
        enrolled = rut in enrolled_ruts
        part = part_by_rut.get(rut, {})
        n_ev = ev_count.get(part.get("id"), 0) if part else 0
        if enrolled:
            estado = "✅ Completada" if part.get("autoevaluacion_completada") else "⏳ Pendiente"
        else:
            estado = ""
        rows_part.append({
            "✓": enrolled,
            "#": i,
            "RUT": rut,
            "Nombres": per.get("pers_nombres", ""),
            "Apellidos": per.get("pers_apellidos", ""),
            "Estado Auto.": estado,
            "N° Eval.": n_ev,
        })

    # --- Planilla tipo Excel ---
    st.markdown(f"#### Participantes ({len(participantes)})")

    if not todas_personas:
        st.info("No hay personas en sist_personas. Crea una usando el formulario de abajo.")
    else:
        df_part = pd.DataFrame(rows_part)
        edited_part = st.data_editor(
            df_part,
            column_config={
                "✓": st.column_config.CheckboxColumn("✓", default=False, width="small"),
                "#": st.column_config.NumberColumn("#", disabled=True, width="small"),
                "RUT": st.column_config.TextColumn("RUT", disabled=True, width=120),
                "Nombres": st.column_config.TextColumn("Nombres", disabled=True, width=160),
                "Apellidos": st.column_config.TextColumn("Apellidos", disabled=True, width=160),
                "Estado Auto.": st.column_config.TextColumn("Estado Auto.", disabled=True, width=140),
                "N° Eval.": st.column_config.NumberColumn("N° Eval.", disabled=True, width="small"),
            },
            disabled=["#", "RUT", "Nombres", "Apellidos", "Estado Auto.", "N° Eval."],
            hide_index=True,
            use_container_width=True,
            key=f"part_editor_{grupo_id}",
        )

        if st.button("Grabar", use_container_width=True, type="primary", key=f"grabar_part_{grupo_id}"):
            nuevos_ruts = [
                todas_personas_ord[idx]["pers_rut"]
                for idx in range(len(rows_part))
                if edited_part.iloc[idx]["✓"] and not rows_part[idx]["✓"]
            ]
            eliminados_ruts = [
                todas_personas_ord[idx]["pers_rut"]
                for idx in range(len(rows_part))
                if not edited_part.iloc[idx]["✓"] and rows_part[idx]["✓"]
            ]
            cambios = False
            errores_elim = []
            for rut in nuevos_ruts:
                queries.crear_participante(grupo_id, rut)
                cambios = True
            for rut in eliminados_ruts:
                part = part_by_rut.get(rut)
                if part:
                    try:
                        queries.eliminar_participante(part["id"])
                        cambios = True
                    except Exception:
                        errores_elim.append(part.get("nombre", rut))
            if nuevos_ruts:
                st.success(f"{len(nuevos_ruts)} participante(s) agregado(s).")
            if eliminados_ruts and not errores_elim:
                st.success(f"{len(eliminados_ruts)} participante(s) eliminado(s).")
            if errores_elim:
                st.error(f"No se pudo eliminar: {', '.join(errores_elim)} (tienen datos vinculados).")
            if not cambios:
                st.info("No hay cambios para guardar.")
            else:
                st.rerun()

    st.divider()

    # --- 3. Botones de proceso (vertical) ---
    st.markdown("#### Acciones del proceso")

    st.markdown("**Invitación a Autoevaluación**")
    st.caption("Envía email a todos los participantes con link a su autoevaluación.")
    if st.button("Enviar Invitación a Autoevaluación", use_container_width=True,
                  disabled=len(participantes) == 0):
        enviados = 0
        errores = 0
        for p in participantes:
            if not p["autoevaluacion_completada"]:
                try:
                    enviar_invitacion_autoevaluacion(p)
                    queries.actualizar_participante(p["id"], {
                        "invitacion_auto_enviada": True,
                    })
                    enviados += 1
                except Exception as e:
                    errores += 1
                    st.error(f"Error enviando a {p['nombre']}: {e}")
        if enviados > 0:
            queries.actualizar_grupo(grupo_id, {"estado": "autoevaluación"})
            st.success(f"{enviados} invitaciones enviadas")
        if errores > 0:
            st.warning(f"{errores} errores")
        if enviados > 0:
            st.rerun()

    st.markdown("")
    st.markdown("**Invitación a Feedback**")
    st.caption("Envía email a evaluadores de participantes que ya completaron su autoevaluación.")
    if st.button("Enviar Invitación a Feedback", use_container_width=True,
                  disabled=len(participantes) == 0):
        enviados = 0
        errores = 0
        for p in participantes:
            if p["autoevaluacion_completada"]:
                evaluadores = queries.listar_evaluadores(p["id"])
                for ev in evaluadores:
                    if not ev["completado"]:
                        try:
                            enviar_invitacion_feedback(ev, p["nombre"])
                            enviados += 1
                        except Exception as e:
                            errores += 1
                            st.error(f"Error enviando a {ev['nombre']}: {e}")
                queries.actualizar_participante(p["id"], {
                    "invitacion_feedback_enviada": True,
                })
        if enviados > 0:
            queries.actualizar_grupo(grupo_id, {"estado": "feedback"})
            st.success(f"{enviados} invitaciones de feedback enviadas")
        if errores > 0:
            st.warning(f"{errores} errores")
        if enviados > 0:
            st.rerun()


# ============================================================
# PÁGINA: SEGUIMIENTO AUTOEVALUACIONES
# ============================================================

def _calcular_estado_auto(participante):
    """Calcula el estado de autoevaluación de un participante."""
    if participante["autoevaluacion_completada"]:
        return "✅ Terminado"
    if not participante["invitacion_auto_enviada"]:
        return "🔲 Sin invitar"
    # Calcular días desde la invitación
    created = participante.get("created_at", "")
    if created:
        try:
            if isinstance(created, str):
                fecha = datetime.fromisoformat(created.replace("Z", "+00:00"))
            else:
                fecha = created
            dias = (datetime.now(timezone.utc) - fecha).days
            if dias > 7:
                return "🔴 Atrasado"
        except Exception:
            pass
    return "🟡 Iniciado"


def _calcular_dias_auto(participante):
    """Calcula días transcurridos desde la invitación."""
    if participante["autoevaluacion_completada"]:
        return "—"
    if not participante["invitacion_auto_enviada"]:
        return "—"
    created = participante.get("created_at", "")
    if created:
        try:
            if isinstance(created, str):
                fecha = datetime.fromisoformat(created.replace("Z", "+00:00"))
            else:
                fecha = created
            dias = (datetime.now(timezone.utc) - fecha).days
            return str(dias)
        except Exception:
            pass
    return "—"


def pagina_seguimiento_auto():
    st.header("Seguimiento Autoevaluaciones")

    grupos = queries.listar_grupos()
    if not grupos:
        st.info("No hay grupos creados.")
        return

    # Recopilar todos los participantes de todos los grupos
    filas = []
    for g in grupos:
        participantes = queries.listar_participantes(g["id"])
        for p in participantes:
            p["_grupo_nombre"] = g["nombre"]
            filas.append(p)

    if not filas:
        st.info("No hay participantes en ningún grupo.")
        return

    # Filtros: Empresa → Grupo
    import pandas as pd

    # Mapa empresa → lista de nombres de grupo
    empresa_a_grupos = {}
    for g in grupos:
        emp = g.get("empresa") or "—"
        empresa_a_grupos.setdefault(emp, []).append(g["nombre"])

    col_emp, col_grp = st.columns([1, 1])
    with col_emp:
        empresas_disponibles = ["Todos"] + sorted(empresa_a_grupos.keys())
        empresa_filtro = st.selectbox("Empresa", options=empresas_disponibles, key="auto_filtro_empresa")
    # Determinar qué grupos corresponden a la empresa seleccionada
    if empresa_filtro == "Todos":
        filas_por_empresa = filas
    else:
        grupos_de_empresa = empresa_a_grupos.get(empresa_filtro, [])
        filas_por_empresa = [p for p in filas if p["_grupo_nombre"] in grupos_de_empresa]

    with col_grp:
        nombres_grupos = sorted({p["_grupo_nombre"] for p in filas_por_empresa})
        grupo_filtro = st.selectbox("Grupo", options=["Todos"] + nombres_grupos, key="auto_filtro_grupo")
    filas_filtradas = filas_por_empresa if grupo_filtro == "Todos" else [p for p in filas_por_empresa if p["_grupo_nombre"] == grupo_filtro]

    st.divider()

    total_auto = len(filas_filtradas)
    term_auto  = sum(1 for p in filas_filtradas if p.get("autoevaluacion_completada"))
    pend_auto  = total_auto - term_auto
    ca, cb, cc = st.columns(3)
    ca.metric("Total", total_auto)
    cb.metric("Terminados", term_auto)
    cc.metric("Pendientes", pend_auto)

    # Construir DataFrame con checkbox de recordatorio (sin columna Grupo)
    part_list = []
    rows = []
    for p in filas_filtradas:
        estado = _calcular_estado_auto(p)
        dias = _calcular_dias_auto(p)
        rows.append({
            "Recordatorio": False,
            "Integrante": p["nombre"],
            "Correo": p["email"],
            "Estado": estado,
            "Días": dias,
            "No Rec.": p.get("num_recordatorios_auto", 0),
        })
        part_list.append(p)

    df = pd.DataFrame(rows)

    edited_df = st.data_editor(
        df,
        column_config={
            "Recordatorio": st.column_config.CheckboxColumn(
                "Recordatorio",
                help="Marca para incluir en el envío de recordatorio",
                default=False,
            ),
            "Integrante": st.column_config.TextColumn("Integrante"),
            "Correo": st.column_config.TextColumn("Correo"),
            "Estado": st.column_config.TextColumn("Estado"),
            "Días": st.column_config.TextColumn("Días"),
            "No Rec.": st.column_config.NumberColumn("No Rec.", format="%d"),
        },
        disabled=["Integrante", "Correo", "Estado", "Días", "No Rec."],
        hide_index=True,
        use_container_width=True,
        key="auto_tabla_editor",
    )

    seleccionados_idx = edited_df.index[edited_df["Recordatorio"] == True].tolist()
    n_sel = len(seleccionados_idx)

    col_btn, _ = st.columns([2, 3])
    with col_btn:
        if st.button(
            f"Enviar Recordatorio ({n_sel} seleccionado{'s' if n_sel != 1 else ''})" if n_sel else "Enviar Recordatorio",
            type="primary",
            use_container_width=True,
            disabled=(n_sel == 0),
        ):
            ok, errores = 0, []
            for idx in seleccionados_idx:
                p = part_list[idx]
                if p["autoevaluacion_completada"]:
                    continue
                try:
                    enviar_recordatorio_autoevaluacion(p)
                    queries.actualizar_participante(p["id"], {
                        "num_recordatorios_auto": (p.get("num_recordatorios_auto") or 0) + 1,
                        "fecha_ultimo_recordatorio_auto": datetime.now(timezone.utc).isoformat(),
                    })
                    ok += 1
                except Exception as e:
                    errores.append(f"{p['nombre']}: {e}")
            if ok:
                st.success(f"✅ {ok} recordatorio(s) enviado(s) correctamente.")
            for err in errores:
                st.error(f"Error: {err}")
            if ok:
                st.rerun()

    # ---- Sección: Informe pendientes ----
    st.divider()
    st.subheader("Informe")

    # Filtrar solo pendientes y ordenar por grupo, luego nombre
    pendientes = [p for p in filas if not p["autoevaluacion_completada"]]
    pendientes.sort(key=lambda p: (p["_grupo_nombre"].lower(), p["nombre"].lower()))

    if not pendientes:
        st.info("No hay autoevaluaciones pendientes.")
    else:
        titulo_informe = "Autoevaluaciones pendientes de responder"
        fecha_hoy = datetime.now(timezone.utc).strftime("%d/%m/%Y")
        encabezados_auto = ["Grupo", "Integrante", "Correo", "Días transcurridos", "Cant. recordatorios"]
        datos_auto = []
        for p in pendientes:
            datos_auto.append([
                p["_grupo_nombre"], p["nombre"], p["email"],
                _calcular_dias_auto(p), str(p.get("num_recordatorios_auto", 0)),
            ])

        if st.button("Generar Informe — Autoevaluaciones pendientes", use_container_width=True):
            st.session_state["mostrar_informe_auto"] = True

        if st.session_state.get("mostrar_informe_auto"):
            col_w, col_x = st.columns(2)
            with col_w:
                # Word
                doc = Document()
                doc.add_heading(titulo_informe, level=1)
                doc.add_paragraph(f"Fecha: {fecha_hoy}")
                table = doc.add_table(rows=1, cols=len(encabezados_auto))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(encabezados_auto):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)
                for fila in datos_auto:
                    row = table.add_row()
                    for i, val in enumerate(fila):
                        row.cells[i].text = val
                        for paragraph in row.cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                buf_w = io.BytesIO()
                doc.save(buf_w)
                buf_w.seek(0)
                st.download_button("Descargar Word", data=buf_w, file_name="autoevaluaciones_pendientes.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="dl_auto_word")
            with col_x:
                # Excel
                wb = Workbook()
                ws = wb.active
                ws.title = "Pendientes"
                header_font = Font(bold=True, color="FFFFFF", size=10)
                header_fill = PatternFill(start_color="1A1A2E", end_color="1A1A2E", fill_type="solid")
                for col_idx, h in enumerate(encabezados_auto, 1):
                    cell = ws.cell(row=1, column=col_idx, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center")
                for row_idx, fila in enumerate(datos_auto, 2):
                    for col_idx, val in enumerate(fila, 1):
                        ws.cell(row=row_idx, column=col_idx, value=val)
                for col in ws.columns:
                    max_len = max(len(str(cell.value or "")) for cell in col)
                    ws.column_dimensions[col[0].column_letter].width = max(max_len + 2, 12)
                buf_x = io.BytesIO()
                wb.save(buf_x)
                buf_x.seek(0)
                st.download_button("Descargar Excel", data=buf_x, file_name="autoevaluaciones_pendientes.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key="dl_auto_excel")


# ============================================================
# PÁGINA: SEGUIMIENTO FEEDBACK
# ============================================================

def _calcular_estado_feedback(evaluador):
    """Calcula el estado de feedback de un evaluador."""
    if evaluador["completado"]:
        return "✅ Terminado"
    created = evaluador.get("created_at", "")
    if created:
        try:
            if isinstance(created, str):
                fecha = datetime.fromisoformat(created.replace("Z", "+00:00"))
            else:
                fecha = created
            dias = (datetime.now(timezone.utc) - fecha).days
            if dias > 7:
                return "🔴 Atrasado"
        except Exception:
            pass
    return "🟡 Iniciado"


def pagina_seguimiento_feedback():
    import pandas as pd
    st.header("Seguimiento Feedback")

    grupos = queries.listar_grupos()
    if not grupos:
        st.info("No hay grupos creados.")
        return

    # Recopilar todas las filas: grupo → participante → evaluador
    filas = []
    for g in grupos:
        participantes = queries.listar_participantes(g["id"])
        for p in participantes:
            if not p["autoevaluacion_completada"]:
                continue
            evaluadores = queries.listar_evaluadores(p["id"])
            for ev in evaluadores:
                filas.append({
                    "grupo_nombre": g["nombre"],
                    "participante_nombre": p["nombre"],
                    "evaluador": ev,
                    "participante": p,
                })

    if not filas:
        st.info("No hay evaluadores registrados aún.")
        return

    # Filtros: Empresa → Grupo
    empresa_a_grupos = {}
    for g in grupos:
        emp = g.get("empresa") or "—"
        empresa_a_grupos.setdefault(emp, []).append(g["nombre"])

    col_emp, col_grp = st.columns([1, 1])
    with col_emp:
        empresas_disponibles = ["Todos"] + sorted(empresa_a_grupos.keys())
        empresa_filtro = st.selectbox("Empresa", options=empresas_disponibles, key="fb_filtro_empresa")
    if empresa_filtro == "Todos":
        filas_por_empresa = filas
    else:
        grupos_de_empresa = empresa_a_grupos.get(empresa_filtro, [])
        filas_por_empresa = [f for f in filas if f["grupo_nombre"] in grupos_de_empresa]

    with col_grp:
        nombres_grupos = sorted({f["grupo_nombre"] for f in filas_por_empresa})
        grupo_filtro = st.selectbox("Grupo", options=["Todos"] + nombres_grupos, key="fb_filtro_grupo")
    filas_filtradas = filas_por_empresa if grupo_filtro == "Todos" else [f for f in filas_por_empresa if f["grupo_nombre"] == grupo_filtro]

    st.divider()

    total_fb = len(filas_filtradas)
    term_fb  = sum(1 for f in filas_filtradas if f["evaluador"].get("completado"))
    pend_fb  = total_fb - term_fb
    fa, fb_, fc = st.columns(3)
    fa.metric("Total", total_fb)
    fb_.metric("Terminados", term_fb)
    fc.metric("Pendientes", pend_fb)

    # Construir DataFrame con checkbox de recordatorio (sin columna Grupo)
    ev_list = []
    rows = []
    for fila in filas_filtradas:
        ev = fila["evaluador"]
        estado = _calcular_estado_feedback(ev)
        rows.append({
            "Recordatorio": False,
            "Evaluador": ev["nombre"],
            "Correo evaluador": ev["email"],
            "Participante": fila["participante_nombre"],
            "Estado": estado,
            "No Rec.": ev.get("num_recordatorios", 0),
        })
        ev_list.append((ev, fila["participante_nombre"]))

    df = pd.DataFrame(rows)

    edited_df = st.data_editor(
        df,
        column_config={
            "Recordatorio": st.column_config.CheckboxColumn(
                "Recordatorio",
                help="Marca para incluir en el envío de recordatorio",
                default=False,
            ),
            "Evaluador": st.column_config.TextColumn("Evaluador"),
            "Correo evaluador": st.column_config.TextColumn("Correo evaluador"),
            "Participante": st.column_config.TextColumn("Participante"),
            "Estado": st.column_config.TextColumn("Estado"),
            "No Rec.": st.column_config.NumberColumn("No Rec.", format="%d"),
        },
        disabled=["Evaluador", "Correo evaluador", "Participante", "Estado", "No Rec."],
        hide_index=True,
        use_container_width=True,
        key="fb_tabla_editor",
    )

    seleccionados_idx = edited_df.index[edited_df["Recordatorio"] == True].tolist()
    n_sel = len(seleccionados_idx)

    col_btn, _ = st.columns([2, 3])
    with col_btn:
        if st.button(
            f"Enviar Recordatorio ({n_sel} seleccionado{'s' if n_sel != 1 else ''})" if n_sel else "Enviar Recordatorio",
            type="primary",
            use_container_width=True,
            disabled=(n_sel == 0),
        ):
            ok, errores = 0, []
            for idx in seleccionados_idx:
                ev, part_nombre = ev_list[idx]
                if ev["completado"]:
                    continue
                try:
                    enviar_recordatorio_feedback(ev, part_nombre)
                    queries.actualizar_evaluador(ev["id"], {
                        "num_recordatorios": (ev.get("num_recordatorios") or 0) + 1,
                        "fecha_ultimo_recordatorio": datetime.now(timezone.utc).isoformat(),
                    })
                    ok += 1
                except Exception as e:
                    errores.append(f"{ev['nombre']}: {e}")
            if ok:
                st.success(f"✅ {ok} recordatorio(s) enviado(s) correctamente.")
            for err in errores:
                st.error(f"Error: {err}")
            if ok:
                st.rerun()

    # ---- Sección: Informe pendientes feedback ----
    st.divider()
    st.subheader("Informe")

    pendientes_fb = [f for f in filas if not f["evaluador"]["completado"]]
    pendientes_fb.sort(key=lambda f: (f["grupo_nombre"].lower(), f["participante_nombre"].lower(), f["evaluador"]["nombre"].lower()))

    if not pendientes_fb:
        st.info("No hay feedback pendientes.")
    else:
        titulo_informe_fb = "Feedback pendientes de responder"
        fecha_hoy_fb = datetime.now(timezone.utc).strftime("%d/%m/%Y")
        encabezados_fb = ["Grupo", "Participante", "Evaluador", "Correo evaluador", "Cant. recordatorios"]
        datos_fb = []
        for f in pendientes_fb:
            ev = f["evaluador"]
            datos_fb.append([
                f["grupo_nombre"], f["participante_nombre"], ev["nombre"],
                ev["email"], str(ev.get("num_recordatorios", 0)),
            ])

        if st.button("Generar Informe — Feedback pendientes", use_container_width=True):
            st.session_state["mostrar_informe_fb"] = True

        if st.session_state.get("mostrar_informe_fb"):
            col_w, col_x = st.columns(2)
            with col_w:
                # Word
                doc = Document()
                doc.add_heading(titulo_informe_fb, level=1)
                doc.add_paragraph(f"Fecha: {fecha_hoy_fb}")
                table = doc.add_table(rows=1, cols=len(encabezados_fb))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(encabezados_fb):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)
                for fila in datos_fb:
                    row = table.add_row()
                    for i, val in enumerate(fila):
                        row.cells[i].text = val
                        for paragraph in row.cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                buf_w = io.BytesIO()
                doc.save(buf_w)
                buf_w.seek(0)
                st.download_button("Descargar Word", data=buf_w, file_name="feedback_pendientes.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="dl_fb_word")
            with col_x:
                # Excel
                wb = Workbook()
                ws = wb.active
                ws.title = "Pendientes"
                header_font = Font(bold=True, color="FFFFFF", size=10)
                header_fill = PatternFill(start_color="1A1A2E", end_color="1A1A2E", fill_type="solid")
                for col_idx, h in enumerate(encabezados_fb, 1):
                    cell = ws.cell(row=1, column=col_idx, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center")
                for row_idx, fila in enumerate(datos_fb, 2):
                    for col_idx, val in enumerate(fila, 1):
                        ws.cell(row=row_idx, column=col_idx, value=val)
                for col in ws.columns:
                    max_len = max(len(str(cell.value or "")) for cell in col)
                    ws.column_dimensions[col[0].column_letter].width = max(max_len + 2, 12)
                buf_x = io.BytesIO()
                wb.save(buf_x)
                buf_x.seek(0)
                st.download_button("Descargar Excel", data=buf_x, file_name="feedback_pendientes.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key="dl_fb_excel")


# ============================================================
# PÁGINA: INFORME FINAL 360
# ============================================================

def _calcular_puntajes_360(participante_id, plantilla_id):
    """Calcula puntajes auto y feedback por competencia y categoría."""
    respuestas = queries.obtener_respuestas_participante(participante_id)
    competencias = queries.listar_competencias_por_plantilla(plantilla_id)

    # Solo feedback de evaluadores que completaron
    evaluadores = queries.listar_evaluadores(participante_id)
    completados_ids = {e["id"] for e in evaluadores if e["completado"]}

    auto_scores = {}
    feedback_scores = {}

    for r in respuestas:
        cid = r["competencia_id"]
        if r["es_autoevaluacion"]:
            auto_scores[cid] = r["puntaje"]
        else:
            ev_id = r.get("evaluador_id")
            if ev_id in completados_ids:
                if cid not in feedback_scores:
                    feedback_scores[cid] = []
                feedback_scores[cid].append((ev_id, r["puntaje"]))

    resultados_comp = []
    for comp in competencias:
        cid = comp["id"]
        auto = auto_scores.get(cid, 0)
        fb_list_raw = feedback_scores.get(cid, [])
        fb_list = [p for _, p in fb_list_raw]
        fb_avg = sum(fb_list) / len(fb_list) if fb_list else 0
        diff = round(fb_avg - auto, 1)
        mejorar = (auto < 3.5) or (fb_avg > 0 and fb_avg < 3.5) or (fb_avg > 0 and fb_avg < auto)
        recomendacion = "Mejorar" if mejorar else "Mantener"
        resultados_comp.append({
            "competencia_id": cid,
            "categoria": comp["categoria_nombre"],
            "texto_auto": comp["texto_auto"],
            "texto_feedback": comp["texto_feedback"],
            "auto": round(auto, 1),
            "feedback": round(fb_avg, 1),
            "diferencia": diff,
            "recomendacion": recomendacion,
            "notas_por_evaluador": {ev_id: p for ev_id, p in fb_list_raw},
        })

    cat_map = {}
    for r in resultados_comp:
        cat = r["categoria"]
        if cat not in cat_map:
            cat_map[cat] = {"auto": [], "feedback": []}
        cat_map[cat]["auto"].append(r["auto"])
        cat_map[cat]["feedback"].append(r["feedback"])

    resultados_cat = []
    for cat, scores in cat_map.items():
        auto_avg = sum(scores["auto"]) / len(scores["auto"])
        fb_avg = sum(scores["feedback"]) / len(scores["feedback"])
        resultados_cat.append({
            "categoria": cat,
            "auto": round(auto_avg, 1),
            "feedback": round(fb_avg, 1),
            "diferencia": round(fb_avg - auto_avg, 1),
        })

    return resultados_cat, resultados_comp


def _generar_contenido_ia(nombre_participante, resultados_cat, resultados_comp):
    """Llama a la API de Anthropic para generar el contenido del informe."""
    # Construir resumen de datos para el prompt
    tabla_cat = "Categoría | Autoevaluación | Feedback | Diferencia\n"
    for c in resultados_cat:
        tabla_cat += f"{c['categoria']} | {c['auto']} | {c['feedback']} | {c['diferencia']}\n"

    tablas_comp = ""
    cats_unicas = []
    for c in resultados_comp:
        if c["categoria"] not in cats_unicas:
            cats_unicas.append(c["categoria"])

    for cat in cats_unicas:
        comps_cat = [c for c in resultados_comp if c["categoria"] == cat]
        tablas_comp += f"\nCategoría: {cat}\n"
        tablas_comp += "Competencia | Auto | Feedback | Diferencia | Recomendación\n"
        for c in comps_cat:
            tablas_comp += f"{c['texto_feedback']} | {c['auto']} | {c['feedback']} | {c['diferencia']} | {c['recomendacion']}\n"

    prompt = f"""Eres un consultor experto en desarrollo organizacional y evaluaciones 360°.
Genera un informe completo en español para {nombre_participante} basándote en estos resultados:

RESULTADOS POR CATEGORÍA:
{tabla_cat}

RESULTADOS POR COMPETENCIA:
{tablas_comp}

La escala es de 1 a 5 (1=Nunca, 2=Rara vez, 3=A veces, 4=Frecuentemente, 5=Siempre).
El umbral de corte es 3.5: sobre ese valor la competencia está consolidada; bajo ese valor requiere desarrollo activo.
"Diferencia" = Feedback - Autoevaluación. Si es positiva, el entorno evalúa mejor de lo que el evaluado se percibe. Si es negativa, existe una sobreestimación propia.
"Mantener" = auto >= 3.5 Y feedback >= 3.5 Y feedback >= auto (competencia consolidada). "Mejorar" = auto < 3.5 O feedback < 3.5 O feedback < auto (área de desarrollo prioritaria).

Genera EXACTAMENTE las siguientes secciones separadas por los marcadores indicados:

===RESUMEN_EJECUTIVO===
Escribe 2-3 párrafos con un resumen ejecutivo de los resultados: fortalezas principales, áreas de mejora, y panorama general.

===ANALISIS_CATEGORIAS===
Escribe un análisis de 1-2 párrafos sobre los resultados a nivel de categorías, destacando las mejor y peor evaluadas.

===PRACTICAS===
Para CADA competencia con recomendación "Mejorar", genera una práctica de aprendizaje con el siguiente formato exacto (una por competencia):

COMPETENCIA: [nombre exacto de la competencia]
OBJETIVO: [objetivo claro y medible]
DESCRIPCIÓN: [descripción de la práctica a realizar, 2-3 oraciones]
PARTICIPANTES: [quiénes deben participar]
DURACIÓN: [tiempo estimado]
KPI: [indicador de medición]
---

===CONCLUSIONES===
Escribe 2-3 párrafos con conclusiones generales y próximos pasos recomendados.

Escribe de forma profesional, directa y constructiva. No uses markdown. Usa lenguaje en tercera persona."""

    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(prompt)
    return response.text


def _parsear_contenido_ia(texto):
    """Parsea el texto de la IA en secciones."""
    secciones = {}
    marcadores = ["===RESUMEN_EJECUTIVO===", "===ANALISIS_CATEGORIAS===", "===PRACTICAS===", "===CONCLUSIONES==="]

    for i, marcador in enumerate(marcadores):
        inicio = texto.find(marcador)
        if inicio == -1:
            continue
        inicio += len(marcador)
        if i + 1 < len(marcadores):
            fin = texto.find(marcadores[i + 1])
            if fin == -1:
                fin = len(texto)
        else:
            fin = len(texto)
        nombre = marcador.replace("===", "").strip()
        secciones[nombre] = texto[inicio:fin].strip()

    return secciones


def _parsear_practicas(texto_practicas):
    """Parsea las prácticas individuales del texto."""
    practicas = []
    bloques = texto_practicas.split("---")
    for bloque in bloques:
        bloque = bloque.strip()
        if not bloque:
            continue
        practica = {}
        for linea in bloque.split("\n"):
            linea = linea.strip()
            for campo in ["COMPETENCIA:", "OBJETIVO:", "DESCRIPCIÓN:", "PARTICIPANTES:", "DURACIÓN:", "KPI:"]:
                if linea.upper().startswith(campo):
                    key = campo.replace(":", "").strip().lower()
                    if key == "descripción":
                        key = "descripcion"
                    if key == "duración":
                        key = "duracion"
                    practica[key] = linea[len(campo):].strip()
        if practica.get("competencia"):
            practicas.append(practica)
    return practicas


def _render_tabla_informe(df, num_cols):
    """Renderiza tabla como HTML con columnas numéricas centradas y coloreadas por signo."""
    import pandas as pd

    def fmt(x):
        if x is None or (isinstance(x, float) and pd.isna(x)):
            return "—"
        if x < 0:
            return f"<b style='color:#cc0000'>{x:.1f}</b>"
        elif x > 0:
            color = "#1a4a7a"
        else:
            color = "inherit"
        return f"<span style='color:{color}'>{x:.1f}</span>"

    fmt_dict = {c: fmt for c in num_cols if c in df.columns}

    html = (
        df.style
        .format(fmt_dict, na_rep="—")
        .set_properties(subset=[c for c in df.columns if c in num_cols],
                        **{"text-align": "center"})
        .set_table_styles([
            {"selector": "th", "props": [
                ("text-align", "center"), ("background-color", "#e8f4f8"),
                ("padding", "6px 14px"), ("font-size", "0.85rem"),
                ("border-bottom", "2px solid #aaa"),
            ]},
            {"selector": "td", "props": [
                ("padding", "4px 14px"), ("font-size", "0.85rem"),
                ("border-bottom", "1px solid #ddd"),
            ]},
            {"selector": "table", "props": [
                ("width", "100%"), ("border-collapse", "collapse"),
            ]},
        ])
        .hide(axis="index")
        .to_html(escape=False)
    )
    st.markdown(html, unsafe_allow_html=True)


def _generar_word_informe(nombre, resultados_cat, resultados_comp, secciones, practicas_parseadas):
    """Genera el documento Word del informe."""
    doc = Document()

    # Título
    titulo = doc.add_heading("INFORME FINAL 360", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Participante: {nombre}")
    doc.add_paragraph(f"Fecha: {datetime.now(timezone.utc).strftime('%d/%m/%Y')}")

    # 1. Resumen Ejecutivo
    doc.add_heading("1. Resumen Ejecutivo", level=1)
    doc.add_paragraph(secciones.get("RESUMEN_EJECUTIVO", ""))

    # 2. Análisis por Categoría
    doc.add_heading("2. Análisis por Categoría", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Categoría", "Autoevaluación", "Feedback", "Diferencia"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for cat in resultados_cat:
        row = table.add_row()
        row.cells[0].text = cat["categoria"]
        row.cells[1].text = f"{cat['auto']:.1f}"
        row.cells[2].text = f"{cat['feedback']:.1f}"
        row.cells[3].text = f"{cat['diferencia']:.1f}"
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_paragraph(secciones.get("ANALISIS_CATEGORIAS", ""))

    # 3. Análisis por Competencia por Categoría
    doc.add_heading("3. Análisis por Competencias por Categoría", level=1)

    cats_unicas = []
    for c in resultados_comp:
        if c["categoria"] not in cats_unicas:
            cats_unicas.append(c["categoria"])

    for cat in cats_unicas:
        doc.add_heading(cat, level=2)
        comps_cat = [c for c in resultados_comp if c["categoria"] == cat]
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        h_comps = ["Competencia", "Autoevaluación", "Feedback", "Diferencia", "Recomendación"]
        for i, h in enumerate(h_comps):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        for comp in comps_cat:
            row = table.add_row()
            row.cells[0].text = comp["texto_feedback"]
            row.cells[1].text = f"{comp['auto']:.1f}"
            row.cells[2].text = f"{comp['feedback']:.1f}"
            row.cells[3].text = f"{comp['diferencia']:.1f}"
            row.cells[4].text = comp["recomendacion"]
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
            if comp["recomendacion"] == "Mejorar":
                run = row.cells[4].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(204, 0, 0)

        # Prácticas para esta categoría
        comps_aprender = [c for c in comps_cat if c["recomendacion"] == "Mejorar"]
        if comps_aprender:
            doc.add_paragraph("")
            p_titulo = doc.add_paragraph()
            run = p_titulo.add_run("Prácticas de aprendizaje:")
            run.bold = True
            run.font.size = Pt(10)

            for comp in comps_aprender:
                practica = None
                for pr in practicas_parseadas:
                    if pr.get("competencia", "").lower() in comp["texto_feedback"].lower() or \
                       comp["texto_feedback"].lower() in pr.get("competencia", "").lower():
                        practica = pr
                        break
                if practica:
                    doc.add_paragraph(f"Competencia: {comp['texto_feedback']}", style="List Bullet")
                    doc.add_paragraph(f"Objetivo: {practica.get('objetivo', '—')}")
                    doc.add_paragraph(f"Descripción: {practica.get('descripcion', '—')}")
                    doc.add_paragraph(f"Participantes: {practica.get('participantes', '—')}")
                    doc.add_paragraph(f"Duración: {practica.get('duracion', '—')}")
                    doc.add_paragraph(f"KPI: {practica.get('kpi', '—')}")
                    doc.add_paragraph("")

    # 4. Conclusiones
    doc.add_heading("4. Conclusiones", level=1)
    doc.add_paragraph(secciones.get("CONCLUSIONES", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def pagina_informe_final():
    st.header("Informe Final 360")

    if not GOOGLE_API_KEY:
        st.warning("Falta configurar la API Key de Google AI en core/config.py (GOOGLE_API_KEY).")
        st.stop()

    grupos = queries.listar_grupos()
    if not grupos:
        st.info("No hay grupos creados.")
        return

    col_emp, col_grp, col_part = st.columns(3)

    empresas = sorted({g.get("empresa") or "—" for g in grupos})
    empresa_sel = col_emp.selectbox("Empresa", options=empresas, key="inf_empresa")

    grupos_filtrados = [g for g in grupos if (g.get("empresa") or "—") == empresa_sel]
    grupo_sel = col_grp.selectbox("Seleccionar grupo", options=grupos_filtrados,
                                   format_func=lambda g: g["nombre"], key="inf_grupo")
    if not grupo_sel:
        return

    participantes = queries.listar_participantes(grupo_sel["id"])
    participantes_completos = sorted(
        [p for p in participantes if p["autoevaluacion_completada"]],
        key=lambda p: (p.get("pers_apellidos") or p.get("nombre") or "").lower(),
    )

    if not participantes_completos:
        col_part.markdown("&nbsp;")
        st.info("No hay participantes con autoevaluación completada en este grupo.")
        return

    part_sel = col_part.selectbox(
        "Seleccionar participante",
        options=participantes_completos,
        format_func=lambda p: p["nombre"],
        key="inf_part",
    )
    if not part_sel:
        return

    # ── Datos base ────────────────────────────────────────────
    grupo_info = queries.obtener_grupo(grupo_sel["id"])
    plantilla_id = grupo_info["plantilla_id"] if grupo_info else None

    st.divider()

    evaluadores = queries.listar_evaluadores(part_sel["id"])
    completados = [e for e in evaluadores if e["completado"]]

    if not completados:
        st.warning("Ningún evaluador ha completado el feedback aún.")
        return

    st.caption(f"Evaluadores completados: {len(completados)}/{len(evaluadores)}")

    if completados and plantilla_id:
        import pandas as pd
        todas_resp = queries.obtener_respuestas_participante(part_sel["id"]) or []
        competencias_pl = queries.listar_competencias_por_plantilla(plantilla_id)
        ev_ids_t     = [e["id"] for e in completados]
        ev_nombres_t = [f"Eva-{i+1}" for i in range(len(completados))]
        filas_ev = []
        for comp in competencias_pl:
            resp_auto = next((r for r in todas_resp if r["competencia_id"] == comp["id"] and r.get("es_autoevaluacion")), None)
            auto_val = resp_auto["puntaje"] if resp_auto else None
            notas_fb = []
            fila = {
                "_cat_orden": comp.get("_cat_orden", 0),
                "_comp_orden": comp.get("orden", 0),
                "Ámbito": comp.get("categoria_nombre", ""),
                "Competencia": comp.get("texto_feedback") or comp.get("nombre", ""),
                "Auto": auto_val,
            }
            for ev_id, ev_nom in zip(ev_ids_t, ev_nombres_t):
                resp = next((r for r in todas_resp if r["competencia_id"] == comp["id"] and r.get("evaluador_id") == ev_id), None)
                nota = resp["puntaje"] if resp else None
                fila[ev_nom] = nota
                if nota is not None:
                    notas_fb.append(nota)
            prom_fb = round(sum(notas_fb) / len(notas_fb), 1) if notas_fb else None
            diff = round(prom_fb - auto_val, 1) if (prom_fb is not None and auto_val is not None) else None
            fila["Prom. Feedback"] = prom_fb
            fila["Diferencia"] = diff
            filas_ev.append(fila)

        # Tabla resumen por ámbito (primero)
        ambitos = {}
        for fila in filas_ev:
            amb = fila["Ámbito"]
            if amb not in ambitos:
                ambitos[amb] = {"auto": [], "prom_fb": [], "ev": {n: [] for n in ev_nombres_t}}
            if fila["Auto"] is not None:
                ambitos[amb]["auto"].append(fila["Auto"])
            if fila["Prom. Feedback"] is not None:
                ambitos[amb]["prom_fb"].append(fila["Prom. Feedback"])
            for ev_nom in ev_nombres_t:
                if fila.get(ev_nom) is not None:
                    ambitos[amb]["ev"][ev_nom].append(fila[ev_nom])

        filas_amb = []
        for amb, vals in ambitos.items():
            auto_amb = round(sum(vals["auto"]) / len(vals["auto"]), 1) if vals["auto"] else None
            prom_amb = round(sum(vals["prom_fb"]) / len(vals["prom_fb"]), 1) if vals["prom_fb"] else None
            diff_amb = round(prom_amb - auto_amb, 1) if (prom_amb is not None and auto_amb is not None) else None
            fila_amb = {"Ámbito": amb, "Auto": auto_amb, "Prom. Feedback": prom_amb, "Diferencia": diff_amb}
            for ev_nom in ev_nombres_t:
                ev_vals = vals["ev"][ev_nom]
                fila_amb[ev_nom] = round(sum(ev_vals) / len(ev_vals), 1) if ev_vals else None
            filas_amb.append(fila_amb)

        cols_amb = ["Ámbito", "Auto", "Prom. Feedback", "Diferencia"] + ev_nombres_t
        df_amb = pd.DataFrame(filas_amb)[cols_amb]
        num_cols_amb = ["Auto", "Prom. Feedback", "Diferencia"] + ev_nombres_t
        st.subheader("Resumen por Ámbito")
        _render_tabla_informe(df_amb, num_cols_amb)

        # Tabla notas por evaluador (después) — ordenada por categoría y luego por competencia
        filas_ev_sorted = sorted(filas_ev, key=lambda f: (f["_cat_orden"], f["_comp_orden"]))
        cols_order = ["Ámbito", "Competencia", "Auto", "Prom. Feedback", "Diferencia"] + ev_nombres_t
        df_ev = pd.DataFrame(filas_ev_sorted)[cols_order]
        num_cols_ev = ["Auto", "Prom. Feedback", "Diferencia"] + ev_nombres_t
        st.subheader("Notas por Evaluador")
        _render_tabla_informe(df_ev, num_cols_ev)
        st.divider()

    if not plantilla_id:
        st.error("No se encontró la plantilla del grupo.")
        return

    if st.button("Generar Informe Final 360", type="primary", use_container_width=True):
        with st.spinner("Calculando puntajes..."):
            resultados_cat, resultados_comp = _calcular_puntajes_360(part_sel["id"], plantilla_id)

        with st.spinner("Generando análisis con IA... (esto puede demorar)"):
            try:
                texto_ia = _generar_contenido_ia(part_sel["nombre"], resultados_cat, resultados_comp)
                secciones = _parsear_contenido_ia(texto_ia)
                practicas_parseadas = _parsear_practicas(secciones.get("PRACTICAS", ""))
                st.session_state["informe_360"] = {
                    "resultados_cat": resultados_cat,
                    "resultados_comp": resultados_comp,
                    "secciones": secciones,
                    "practicas": practicas_parseadas,
                    "nombre": part_sel["nombre"],
                    "evaluadores": completados,
                }
            except Exception as e:
                st.error(f"Error al generar informe: {e}")
                return

    # Mostrar informe si ya fue generado
    if "informe_360" not in st.session_state:
        return

    inf = st.session_state["informe_360"]
    resultados_cat = inf["resultados_cat"]
    resultados_comp = inf["resultados_comp"]
    secciones = inf["secciones"]
    practicas_parseadas = inf["practicas"]
    nombre = inf["nombre"]

    st.divider()
    st.markdown(f"## INFORME FINAL 360")
    st.markdown(f"**Participante:** {nombre}")

    # 1. Resumen Ejecutivo
    st.markdown("### 1. Resumen Ejecutivo")
    st.markdown(secciones.get("RESUMEN_EJECUTIVO", ""))

    # 2. Análisis por Categoría
    st.markdown("### 2. Análisis por Categoría")
    hc = st.columns([3, 1.5, 1.5, 1.5])
    hc[0].markdown("**Categoría**")
    hc[1].markdown("**Autoevaluación**")
    hc[2].markdown("**Feedback**")
    hc[3].markdown("**Diferencia**")
    st.markdown("---")
    def _fmt_val(col, v):
        if v < 0:
            col.markdown(f"**:red[{v:.1f}]**")
        else:
            col.caption(f"{v:.1f}")

    for cat in resultados_cat:
        rc = st.columns([3, 1.5, 1.5, 1.5])
        rc[0].caption(cat["categoria"])
        _fmt_val(rc[1], cat["auto"])
        _fmt_val(rc[2], cat["feedback"])
        _fmt_val(rc[3], cat["diferencia"])

    st.markdown("")
    st.markdown(secciones.get("ANALISIS_CATEGORIAS", ""))

    # 3. Análisis por Competencia
    st.markdown("### 3. Análisis por Competencias por Categoría")

    cats_unicas = []
    for c in resultados_comp:
        if c["categoria"] not in cats_unicas:
            cats_unicas.append(c["categoria"])

    for cat in cats_unicas:
        st.markdown(f"#### {cat}")
        comps_cat = [c for c in resultados_comp if c["categoria"] == cat]

        hc2 = st.columns([3, 1.2, 1.2, 1.2, 1.5])
        hc2[0].markdown("**Competencia**")
        hc2[1].markdown("**Auto**")
        hc2[2].markdown("**Feedback**")
        hc2[3].markdown("**Diferencia**")
        hc2[4].markdown("**Recomendación**")
        st.markdown("---")

        for comp in comps_cat:
            rc2 = st.columns([3, 1.2, 1.2, 1.2, 1.5])
            rc2[0].caption(comp["texto_feedback"])
            _fmt_val(rc2[1], comp["auto"])
            _fmt_val(rc2[2], comp["feedback"])
            _fmt_val(rc2[3], comp["diferencia"])
            if comp["recomendacion"] == "Mejorar":
                rc2[4].markdown(f"**:red[{comp['recomendacion']}]**")
            else:
                rc2[4].caption(comp["recomendacion"])

        # Prácticas
        comps_aprender = [c for c in comps_cat if c["recomendacion"] == "Mejorar"]
        if comps_aprender:
            st.markdown("**Prácticas de aprendizaje:**")
            for comp in comps_aprender:
                practica = None
                for pr in practicas_parseadas:
                    if pr.get("competencia", "").lower() in comp["texto_feedback"].lower() or \
                       comp["texto_feedback"].lower() in pr.get("competencia", "").lower():
                        practica = pr
                        break
                if practica:
                    with st.container(border=True):
                        st.markdown(f"**{comp['texto_feedback']}**")
                        st.caption(f"Objetivo: {practica.get('objetivo', '—')}")
                        st.caption(f"Descripción: {practica.get('descripcion', '—')}")
                        st.caption(f"Participantes: {practica.get('participantes', '—')}")
                        st.caption(f"Duración: {practica.get('duracion', '—')}")
                        st.caption(f"KPI: {practica.get('kpi', '—')}")

        st.divider()

    # 4. Conclusiones
    st.markdown("### 4. Conclusiones")
    st.markdown(secciones.get("CONCLUSIONES", ""))

    # Botones de descarga
    st.divider()
    st.subheader("Descargar informe")
    col_w, col_x = st.columns(2)

    with col_w:
        buf_word = _generar_word_informe(nombre, resultados_cat, resultados_comp, secciones, practicas_parseadas)
        st.download_button("Descargar Word", data=buf_word,
                            file_name=f"informe_360_{nombre.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_inf360_word")

    with col_x:
        # Excel con dos hojas
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Por Categoría"
        h_font = Font(bold=True, color="FFFFFF", size=10)
        h_fill = PatternFill(start_color="1A1A2E", end_color="1A1A2E", fill_type="solid")
        for i, h in enumerate(["Categoría", "Autoevaluación", "Feedback", "Diferencia"], 1):
            cell = ws1.cell(row=1, column=i, value=h)
            cell.font = h_font
            cell.fill = h_fill
        for idx, cat in enumerate(resultados_cat, 2):
            ws1.cell(row=idx, column=1, value=cat["categoria"])
            for col_i, key in [(2, "auto"), (3, "feedback"), (4, "diferencia")]:
                c = ws1.cell(row=idx, column=col_i, value=cat[key])
                c.number_format = "0.0"
        for col in ws1.columns:
            max_len = max(len(str(c.value or "")) for c in col)
            ws1.column_dimensions[col[0].column_letter].width = max(max_len + 2, 14)

        ws2 = wb.create_sheet("Por Competencia")
        ev_list = inf.get("evaluadores", [])
        ev_ids = [e["id"] for e in ev_list]
        ev_nombres = [e["nombre"] for e in ev_list]
        headers2 = ["Categoría", "Competencia", "Autoevaluación"] + ev_nombres + ["Promedio Feedback", "Diferencia", "Recomendación"]
        for i, h in enumerate(headers2, 1):
            cell = ws2.cell(row=1, column=i, value=h)
            cell.font = h_font
            cell.fill = h_fill
        for idx, comp in enumerate(resultados_comp, 2):
            col = 1
            ws2.cell(row=idx, column=col, value=comp["categoria"]); col += 1
            ws2.cell(row=idx, column=col, value=comp["texto_feedback"]); col += 1
            c = ws2.cell(row=idx, column=col, value=comp["auto"]); c.number_format = "0.0"; col += 1
            for ev_id in ev_ids:
                nota = comp.get("notas_por_evaluador", {}).get(ev_id, "")
                c2 = ws2.cell(row=idx, column=col, value=nota)
                if nota != "":
                    c2.number_format = "0.0"
                col += 1
            c = ws2.cell(row=idx, column=col, value=comp["feedback"]); c.number_format = "0.0"; col += 1
            c = ws2.cell(row=idx, column=col, value=comp["diferencia"]); c.number_format = "0.0"; col += 1
            ws2.cell(row=idx, column=col, value=comp["recomendacion"]); col += 1
        for col in ws2.columns:
            max_len = max(len(str(c.value or "")) for c in col)
            ws2.column_dimensions[col[0].column_letter].width = max(max_len + 2, 14)

        buf_x = io.BytesIO()
        wb.save(buf_x)
        buf_x.seek(0)
        st.download_button("Descargar Excel", data=buf_x,
                            file_name=f"informe_360_{nombre.replace(' ', '_')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="dl_inf360_excel")


# ============================================================
# PÁGINA: IMPORTAR ENCUESTA DESDE CSV
# ============================================================

def pagina_plantillas():
    st.header("Ingreso Encuestas")

    if "plantilla_editando" not in st.session_state:
        st.session_state["plantilla_editando"] = None
    if "editando_comp" not in st.session_state:
        st.session_state["editando_comp"] = None
    if "plantilla_vista" not in st.session_state:
        st.session_state["plantilla_vista"] = None
    if "np_categorias" not in st.session_state:
        st.session_state["np_categorias"] = []
    if "np_competencias" not in st.session_state:
        st.session_state["np_competencias"] = []

    if st.session_state["plantilla_editando"]:
        _editar_plantilla(st.session_state["plantilla_editando"])
        return

    plantillas = queries.listar_plantillas()
    nombres_existentes = [p["nombre"].lower().strip() for p in plantillas]

    tab_lista, tab_nueva, tab_csv = st.tabs(["Plantillas existentes", "Agregar nueva", "Importar desde CSV"])

    # ---- Pestaña: Plantillas existentes ----
    with tab_lista:
        if not plantillas:
            st.info("No hay plantillas creadas.")
        else:
            hc = st.columns([3, 1.5, 1.5, 0.8, 0.8, 0.8])
            hc[0].markdown("**Nombre**")
            hc[1].markdown("**Categorías**")
            hc[2].markdown("**Competencias**")
            hc[3].markdown("")
            hc[4].markdown("")
            hc[5].markdown("")
            st.markdown("---")

            for p in plantillas:
                cats = queries.listar_categorias(p["id"])
                n_comps = sum(len(queries.listar_competencias(c["id"])) for c in cats)
                rc = st.columns([3, 1.5, 1.5, 0.8, 0.8, 0.8])
                with rc[0]:
                    st.markdown(f"**{p['nombre']}**")
                with rc[1]:
                    st.markdown(str(len(cats)))
                with rc[2]:
                    st.markdown(str(n_comps))
                with rc[3]:
                    if st.button("Ver", key=f"ver_{p['id']}", use_container_width=True):
                        if st.session_state.get("plantilla_vista") == p["id"]:
                            st.session_state["plantilla_vista"] = None
                        else:
                            st.session_state["plantilla_vista"] = p["id"]
                        st.rerun()
                with rc[4]:
                    if st.button("Editar", key=f"edit_{p['id']}", use_container_width=True):
                        st.session_state["plantilla_editando"] = p["id"]
                        st.rerun()
                with rc[5]:
                    if st.button("Eliminar", key=f"del_{p['id']}", use_container_width=True):
                        n_grupos = queries.contar_grupos_por_plantilla(p["id"])
                        if n_grupos > 0:
                            st.error(
                                f"No se puede eliminar: esta plantilla tiene {n_grupos} "
                                f"grupo{'s' if n_grupos > 1 else ''} asociado{'s' if n_grupos > 1 else ''}. "
                                "Elimina los grupos primero."
                            )
                        else:
                            queries.eliminar_plantilla(p["id"])
                            st.rerun()

                if st.session_state.get("plantilla_vista") == p["id"]:
                    _vista_previa_plantilla(p["id"], cats)

    # ---- Pestaña: Agregar nueva ----
    with tab_nueva:
        st.subheader("Nueva plantilla")

        nombre_np = st.text_input("Nombre de la plantilla", key="np_nombre")
        descripcion_np = st.text_input("Descripción (opcional)", key="np_descripcion")

        st.divider()

        st.markdown("**Categorías**")
        if st.session_state["np_categorias"]:
            for i, cat_name in enumerate(st.session_state["np_categorias"]):
                cc1, cc2 = st.columns([5, 1])
                with cc1:
                    st.markdown(f"{i + 1}. {cat_name}")
                with cc2:
                    if st.button("Quitar", key=f"np_delcat_{i}", use_container_width=True):
                        st.session_state["np_competencias"] = [
                            c for c in st.session_state["np_competencias"] if c["categoria"] != cat_name
                        ]
                        st.session_state["np_categorias"].pop(i)
                        st.rerun()
        else:
            st.caption("Sin categorías aún.")

        nc1, nc2 = st.columns([4, 1])
        with nc1:
            nueva_cat_nombre = st.text_input("Nueva categoría", key="np_nueva_cat", placeholder="Nombre de categoría")
        with nc2:
            st.markdown("")
            if st.button("Agregar categoría", use_container_width=True):
                if nueva_cat_nombre.strip():
                    if nueva_cat_nombre.strip() in st.session_state["np_categorias"]:
                        st.warning("Esa categoría ya existe.")
                    else:
                        st.session_state["np_categorias"].append(nueva_cat_nombre.strip())
                        st.session_state["np_nueva_cat"] = ""
                        st.rerun()
                else:
                    st.warning("Ingresa un nombre de categoría.")

        st.divider()

        st.markdown("**Competencias**")
        if st.session_state["np_competencias"]:
            hcomp = st.columns([1.5, 2.5, 2.5, 0.8])
            hcomp[0].markdown("**Categoría**")
            hcomp[1].markdown("**Texto auto (1ª persona)**")
            hcomp[2].markdown("**Texto feedback (3ª persona)**")
            hcomp[3].markdown("")
            st.markdown("---")
            for i, comp in enumerate(st.session_state["np_competencias"]):
                rcomp = st.columns([1.5, 2.5, 2.5, 0.8])
                with rcomp[0]:
                    st.caption(comp["categoria"])
                with rcomp[1]:
                    st.markdown(comp["texto_auto"])
                with rcomp[2]:
                    st.markdown(comp["texto_feedback"])
                with rcomp[3]:
                    if st.button("Quitar", key=f"np_delcomp_{i}", use_container_width=True):
                        st.session_state["np_competencias"].pop(i)
                        st.rerun()
        else:
            st.caption("Sin competencias aún.")

        if not st.session_state["np_categorias"]:
            st.info("Agrega al menos una categoría para poder agregar competencias.")
        else:
            with st.form("np_nueva_comp"):
                fc1, fc2, fc3, fc4 = st.columns([1.5, 2.5, 2.5, 1])
                with fc1:
                    cat_comp_sel = st.selectbox(
                        "Categoría", options=st.session_state["np_categorias"], label_visibility="collapsed"
                    )
                with fc2:
                    txt_auto_np = st.text_area(
                        "Texto auto", placeholder="Ej: Comunico mis ideas con claridad",
                        label_visibility="collapsed", height=80, key="np_txt_auto"
                    )
                with fc3:
                    txt_feed_np = st.text_area(
                        "Texto feedback", placeholder="Ej: Comunica sus ideas con claridad",
                        label_visibility="collapsed", height=80, key="np_txt_feed"
                    )
                with fc4:
                    st.markdown("")
                    agregar_comp = st.form_submit_button("Agregar competencia")
                if agregar_comp:
                    if txt_auto_np.strip() and txt_feed_np.strip():
                        st.session_state["np_competencias"].append({
                            "categoria": cat_comp_sel,
                            "texto_auto": txt_auto_np.strip(),
                            "texto_feedback": txt_feed_np.strip(),
                        })
                        st.session_state["np_txt_auto"] = ""
                        st.session_state["np_txt_feed"] = ""
                        st.rerun()
                    else:
                        st.warning("Completa ambos textos.")

        st.divider()

        if st.button("Crear Plantilla", use_container_width=True, type="primary"):
            errores = []
            if not nombre_np.strip():
                errores.append("Ingresa un nombre para la plantilla.")
            elif nombre_np.strip().lower() in nombres_existentes:
                errores.append(f"Ya existe una plantilla con el nombre '{nombre_np.strip()}'.")
            if not st.session_state["np_categorias"]:
                errores.append("Agrega al menos una categoría.")
            if not st.session_state["np_competencias"]:
                errores.append("Agrega al menos una competencia.")

            if errores:
                for err in errores:
                    st.error(err)
            else:
                try:
                    plantilla = queries.crear_plantilla(nombre_np.strip(), descripcion_np.strip())
                    cat_map = {}
                    for idx, cat_name in enumerate(st.session_state["np_categorias"]):
                        cat_obj = queries.crear_categoria(plantilla["id"], cat_name, idx + 1)
                        cat_map[cat_name] = cat_obj["id"]
                    conteo_cat = {}
                    for comp in st.session_state["np_competencias"]:
                        cat_id = cat_map[comp["categoria"]]
                        conteo_cat[cat_id] = conteo_cat.get(cat_id, 0) + 1
                        queries.crear_competencia(cat_id, comp["texto_auto"], comp["texto_feedback"], conteo_cat[cat_id])
                    st.session_state["np_categorias"] = []
                    st.session_state["np_competencias"] = []
                    st.session_state["np_nombre"] = ""
                    st.session_state["np_descripcion"] = ""
                    st.session_state["np_nueva_cat"] = ""
                    st.session_state["np_txt_auto"] = ""
                    st.session_state["np_txt_feed"] = ""
                    st.success("Plantilla creada exitosamente.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al crear: {e}")

    # ---- Pestaña: Importar desde CSV ----
    with tab_csv:
        _contenido_importar_encuesta_csv()


def _contenido_importar_encuesta_csv():
    import pandas as pd
    st.markdown(
        "Carga un CSV con **Ámbito** y **Competencia** (en primera persona). "
        "La IA genera automáticamente la versión de feedback en tercera persona singular."
    )

    # ---- Formato esperado ----
    with st.expander("Ver formato esperado del CSV"):
        ej = pd.DataFrame({
            "Ámbito":      ["Liderazgo", "Liderazgo", "Comunicación"],
            "Competencia": [
                "Inspiro y motivo a mi equipo hacia los objetivos",
                "Comunico con claridad mis expectativas",
                "Escucho activamente a mis compañeros",
            ],
        })
        st.dataframe(ej, use_container_width=True, hide_index=True)
        buf_ej = io.BytesIO()
        ej.to_csv(buf_ej, index=False, sep=";", encoding="utf-8-sig")
        buf_ej.seek(0)
        st.download_button(
            "⬇️ Descargar plantilla CSV",
            data=buf_ej,
            file_name="plantilla_encuesta_360.csv",
            mime="text/csv",
        )

    st.divider()

    # ---- Subida del archivo ----
    archivo = st.file_uploader("Seleccionar archivo CSV", type=["csv"])
    if archivo is None:
        st.info("Sube un archivo CSV para continuar.")
        return

    try:
        df = pd.read_csv(archivo, sep=None, engine="python", encoding="utf-8-sig")
    except Exception:
        archivo.seek(0)
        try:
            df = pd.read_csv(archivo, sep=None, engine="python", encoding="latin1")
        except Exception as e:
            st.error(f"No se pudo leer el archivo: {e}")
            return

    df.columns = df.columns.str.strip()

    # ---- Mapeo de columnas ----
    st.markdown("#### Mapeo de columnas")
    cols = list(df.columns)

    def _detectar(opciones, palabras):
        for c in opciones:
            if any(p in c.lower() for p in palabras):
                return c
        return opciones[0]

    default_amb  = _detectar(cols, ["ámbito", "ambito", "categoria", "categoría", "dimension", "área", "area"])
    default_comp = _detectar(cols, ["competencia", "texto", "indicador", "pregunta", "descripcion"])

    mc1, mc2 = st.columns(2)
    col_amb  = mc1.selectbox("Columna Ámbito", cols, index=cols.index(default_amb))
    col_comp = mc2.selectbox("Columna Competencia (1ª persona)", cols, index=cols.index(default_comp))

    if col_amb == col_comp:
        st.warning("Las columnas seleccionadas son la misma. Elige columnas distintas.")
        return

    # Limpiar datos
    df_v = df[[col_amb, col_comp]].dropna().copy()
    df_v[col_amb]  = df_v[col_amb].astype(str).str.strip()
    df_v[col_comp] = df_v[col_comp].astype(str).str.strip()
    df_v = df_v[(df_v[col_amb] != "") & (df_v[col_comp] != "")]

    if df_v.empty:
        st.error("No se encontraron filas válidas con las columnas seleccionadas.")
        return

    ambitos_unicos = list(dict.fromkeys(df_v[col_amb].tolist()))
    pm1, pm2 = st.columns(2)
    pm1.metric("Ámbitos / Categorías", len(ambitos_unicos))
    pm2.metric("Competencias", len(df_v))

    with st.expander("Ver competencias detectadas"):
        for amb in ambitos_unicos:
            comps = df_v[df_v[col_amb] == amb][col_comp].tolist()
            st.markdown(
                f"<div style='background:#1a3a5c;color:white;padding:5px 12px;"
                f"border-radius:5px;margin:8px 0 4px;font-weight:bold'>{amb}</div>",
                unsafe_allow_html=True,
            )
            for i, c in enumerate(comps, 1):
                st.markdown(f"&nbsp;&nbsp;&nbsp;{i}. {c}")

    st.divider()

    # ---- Generar versión feedback con IA ----
    st.markdown("#### Generar versión feedback (3ª persona singular)")
    st.caption(
        "La IA convertirá cada competencia de primera persona a tercera persona singular. "
        "Podrás revisar y editar el resultado antes de guardar."
    )

    if st.button("🤖 Generar versión feedback con IA", use_container_width=True, type="primary"):
        lista_numerada = "\n".join(
            f"{i+1}. {row[col_comp]}"
            for i, (_, row) in enumerate(df_v.iterrows())
        )
        prompt = (
            "Eres un experto en evaluaciones de desempeño 360°.\n"
            "Tienes una lista de afirmaciones escritas en PRIMERA PERSONA SINGULAR "
            "(autoevaluación). Conviértelas a TERCERA PERSONA SINGULAR para su uso "
            "como preguntas de feedback de evaluadores hacia el evaluado.\n\n"
            "Reglas estrictas:\n"
            "- Mantén el mismo significado y nivel de detalle\n"
            "- Solo cambia la persona gramatical (yo → él/ella)\n"
            "- No agregues el nombre del sujeto, usa frases impersonales o 'el/la evaluado/a'\n"
            "- Devuelve SOLO la lista numerada con las frases convertidas, sin explicaciones ni encabezados\n"
            "- Un ítem por línea, en el mismo orden\n\n"
            f"Lista ({len(df_v)} ítems):\n{lista_numerada}"
        )
        with st.spinner("Generando versión feedback..."):
            try:
                genai.configure(api_key=GOOGLE_API_KEY)
                model = genai.GenerativeModel("gemini-2.5-flash")
                respuesta = model.generate_content(prompt)
                lineas = [
                    l.strip()
                    for l in respuesta.text.strip().split("\n")
                    if l.strip()
                ]
                # Quitar numeración si la IA la incluyó
                import re
                feedback_textos = [
                    re.sub(r'^\d+[\.\)]\s*', '', l).strip()
                    for l in lineas
                    if l.strip()
                ]
                if len(feedback_textos) == len(df_v):
                    st.session_state["imp360_feedback"] = feedback_textos
                    st.session_state["imp360_df"] = df_v.reset_index(drop=True)
                    st.session_state["imp360_col_amb"] = col_amb
                    st.session_state["imp360_col_comp"] = col_comp
                else:
                    st.warning(
                        f"La IA devolvió {len(feedback_textos)} frases pero se esperaban {len(df_v)}. "
                        "Intenta de nuevo o ajusta el CSV."
                    )
            except Exception as e:
                st.error(f"Error al contactar la IA: {e}")

    # ---- Revisión y edición del resultado ----
    if (
        "imp360_feedback" in st.session_state
        and "imp360_df" in st.session_state
        and len(st.session_state["imp360_feedback"]) == len(st.session_state["imp360_df"])
    ):
        df_imp   = st.session_state["imp360_df"]
        col_amb_ = st.session_state["imp360_col_amb"]
        col_comp_= st.session_state["imp360_col_comp"]
        feedbacks= st.session_state["imp360_feedback"]

        st.divider()
        st.markdown("#### Revisión — Autoevaluación vs Feedback")
        st.caption("Puedes editar los textos de feedback antes de guardar.")

        feedbacks_editados = []
        for amb in list(dict.fromkeys(df_imp[col_amb_].tolist())):
            st.markdown(
                f"<div style='background:#1a3a5c;color:white;padding:5px 12px;"
                f"border-radius:5px;margin:12px 0 4px;font-weight:bold'>{amb}</div>",
                unsafe_allow_html=True,
            )
            indices = df_imp[df_imp[col_amb_] == amb].index.tolist()
            for idx in indices:
                texto_auto = df_imp.loc[idx, col_comp_]
                texto_fb_orig = feedbacks[idx]
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown(
                        f"<div style='background:#f0f4f8;border-radius:5px;padding:8px 10px;"
                        f"font-size:0.88rem;color:#333'>"
                        f"<span style='color:#666;font-size:0.75rem'>AUTOEVALUACIÓN (1ª persona)</span><br>"
                        f"{texto_auto}</div>",
                        unsafe_allow_html=True,
                    )
                with c2:
                    texto_fb_edit = st.text_area(
                        "Feedback (3ª persona)",
                        value=texto_fb_orig,
                        key=f"fb_edit_{idx}",
                        height=80,
                        label_visibility="collapsed",
                    )
                feedbacks_editados.append((amb, texto_auto, texto_fb_edit))

        st.divider()

        # ---- Datos de la nueva plantilla ----
        st.markdown("#### Datos de la nueva Plantilla")
        nombre_pl = st.text_input("Nombre de la plantilla", placeholder="Ej: Evaluación 360° Liderazgo 2025")
        desc_pl   = st.text_area("Descripción (opcional)", height=70)

        if st.button(
            "💾 Guardar plantilla",
            type="primary",
            use_container_width=True,
            disabled=not nombre_pl.strip(),
        ):
            with st.spinner("Guardando..."):
                try:
                    plantilla = queries.crear_plantilla(nombre_pl.strip(), desc_pl.strip())
                    pid = plantilla["id"]

                    orden_cat = 1
                    cat_id_map = {}
                    conteo_comp = {}
                    ambitos_orden = list(dict.fromkeys(
                        [t[0] for t in feedbacks_editados]
                    ))
                    for amb_nombre in ambitos_orden:
                        cat_obj = queries.crear_categoria(pid, amb_nombre, orden_cat)
                        cat_id_map[amb_nombre] = cat_obj["id"]
                        conteo_comp[cat_obj["id"]] = 1
                        orden_cat += 1

                    for amb_nombre, texto_auto, texto_fb in feedbacks_editados:
                        cid = cat_id_map[amb_nombre]
                        queries.crear_competencia(cid, texto_auto, texto_fb, conteo_comp[cid])
                        conteo_comp[cid] += 1

                    n_cats  = len(ambitos_orden)
                    n_comps = len(feedbacks_editados)
                    st.success(
                        f"✅ Plantilla **\"{nombre_pl.strip()}\"** creada con "
                        f"**{n_cats} categoría(s)** y **{n_comps} competencia(s)**."
                    )
                    # Limpiar estado
                    for k in ["imp360_feedback", "imp360_df", "imp360_col_amb", "imp360_col_comp"]:
                        st.session_state.pop(k, None)
                    st.balloons()
                except Exception as e:
                    st.error(f"Error al guardar: {e}")


# ============================================================
# CONTENIDO: IMPORTAR PARTICIPANTES DESDE CSV
# ============================================================

def _contenido_importar_participantes():
    st.info(
        "La importación desde CSV ya no está disponible. "
        "Los participantes se agregan seleccionándolos directamente desde el registro de personas "
        "del Sistema de Asistencias (sist_personas). "
        "Si la persona no existe, créala usando el formulario **'Crear persona en Asistencias'** "
        "en el detalle del grupo."
    )
    return

    import pandas as pd  # unreachable — kept for reference

    st.markdown(
        "Carga un CSV con **Nombre** y **Correo** de los participantes. "
        "Se agregarán al grupo que selecciones, omitiendo duplicados."
    )

    # ---- Formato esperado ----
    with st.expander("Ver formato esperado del CSV"):
        ej = pd.DataFrame({
            "Nombre": ["Juan Pérez", "María González", "Carlos López"],
            "Correo": ["juan@empresa.com", "maria@empresa.com", "carlos@empresa.com"],
        })
        st.dataframe(ej, use_container_width=True, hide_index=True)
        buf_ej = io.BytesIO()
        ej.to_csv(buf_ej, index=False, sep=";", encoding="utf-8-sig")
        buf_ej.seek(0)
        st.download_button(
            "⬇️ Descargar plantilla CSV",
            data=buf_ej,
            file_name="plantilla_participantes_360.csv",
            mime="text/csv",
        )

    st.divider()

    # ---- Selección empresa → grupo ----
    grupos = queries.listar_grupos()
    if not grupos:
        st.warning("No hay grupos creados. Crea uno primero en Ingreso de Grupos.")
        return

    empresas = sorted({g.get("empresa") or "—" for g in grupos})
    empresa_sel = st.selectbox("Empresa", options=empresas, key="imp_part_empresa")

    grupos_filtrados = [g for g in grupos if (g.get("empresa") or "—") == empresa_sel]
    if not grupos_filtrados:
        st.info("No hay grupos para esta empresa.")
        return

    grupo_sel = st.selectbox(
        "Grupo destino",
        options=grupos_filtrados,
        format_func=lambda g: g["nombre"],
        key="imp_part_grupo",
    )

    st.divider()

    # ---- Subida del archivo ----
    archivo = st.file_uploader("Seleccionar archivo CSV", type=["csv"], key="imp_part_file")
    if archivo is None:
        st.info("Sube un archivo CSV para continuar.")
        return

    try:
        df = pd.read_csv(archivo, sep=None, engine="python", encoding="utf-8-sig")
    except Exception:
        archivo.seek(0)
        try:
            df = pd.read_csv(archivo, sep=None, engine="python", encoding="latin1")
        except Exception as e:
            st.error(f"No se pudo leer el archivo: {e}")
            return

    df.columns = df.columns.str.strip()

    # ---- Mapeo de columnas ----
    st.markdown("#### Mapeo de columnas")
    cols = list(df.columns)

    def _detectar(opciones, palabras):
        for c in opciones:
            if any(p in c.lower() for p in palabras):
                return c
        return opciones[0]

    default_nom   = _detectar(cols, ["nombre", "name", "participante", "integrante"])
    default_email = _detectar(cols, ["correo", "email", "mail", "e-mail"])

    mc1, mc2 = st.columns(2)
    col_nom   = mc1.selectbox("Columna Nombre", cols, index=cols.index(default_nom), key="imp_part_col_nom")
    col_email = mc2.selectbox("Columna Correo", cols, index=cols.index(default_email), key="imp_part_col_email")

    if col_nom == col_email:
        st.warning("Las columnas seleccionadas son la misma. Elige columnas distintas.")
        return

    # Limpiar datos
    df_v = df[[col_nom, col_email]].dropna().copy()
    df_v[col_nom]   = df_v[col_nom].astype(str).str.strip()
    df_v[col_email] = df_v[col_email].astype(str).str.strip()
    df_v = df_v[(df_v[col_nom] != "") & (df_v[col_email] != "")]

    if df_v.empty:
        st.error("No se encontraron filas válidas con las columnas seleccionadas.")
        return

    # Detectar duplicados con participantes ya existentes en el grupo
    existentes = queries.listar_participantes(grupo_sel["id"])
    emails_existentes = {p["email"].strip().lower() for p in existentes}

    df_nuevos = df_v[~df_v[col_email].str.lower().isin(emails_existentes)].reset_index(drop=True)
    df_dupes  = df_v[df_v[col_email].str.lower().isin(emails_existentes)].reset_index(drop=True)

    m1, m2, m3 = st.columns(3)
    m1.metric("Total en CSV", len(df_v))
    m2.metric("Nuevos a importar", len(df_nuevos))
    m3.metric("Ya existentes (omitir)", len(df_dupes))

    if df_dupes.empty is False:
        with st.expander(f"Ver {len(df_dupes)} participante(s) ya registrado(s) — serán omitidos"):
            st.dataframe(df_dupes.rename(columns={col_nom: "Nombre", col_email: "Correo"}),
                         use_container_width=True, hide_index=True)

    if df_nuevos.empty:
        st.info("Todos los participantes del CSV ya están registrados en este grupo.")
        return

    # Vista previa de los que se importarán
    st.markdown("#### Vista previa — Participantes a importar")
    st.dataframe(
        df_nuevos.rename(columns={col_nom: "Nombre", col_email: "Correo"}),
        use_container_width=True,
        hide_index=True,
    )

    st.divider()

    if st.button(
        f"Importar {len(df_nuevos)} participante(s) al grupo \"{grupo_sel['nombre']}\"",
        type="primary",
        use_container_width=True,
    ):
        ok = 0
        errores = []
        for _, row in df_nuevos.iterrows():
            try:
                queries.crear_participante(grupo_sel["id"], row[col_nom], row[col_email])
                ok += 1
            except Exception as e:
                errores.append(f"{row[col_nom]} — {e}")

        if ok:
            st.success(f"✅ {ok} participante(s) importado(s) correctamente al grupo \"{grupo_sel['nombre']}\".")
        if errores:
            for err in errores:
                st.error(f"Error: {err}")


# ============================================================
# INGRESOS ESPECIALES — helpers de cascada y tabs
# ============================================================

_ESCALA_LABELS = {
    1: "1 – Nunca",
    2: "2 – Rara vez",
    3: "3 – A veces",
    4: "4 – Frecuentemente",
    5: "5 – Siempre",
}


def _cascada_ingresos(prefix):
    """Renderiza selectores Empresa → Grupo → Participante | Correo.
    Devuelve (grupo_sel, part_sel) o (None, None) si falta alguno."""
    grupos = queries.listar_grupos()
    if not grupos:
        st.warning("No hay grupos creados.")
        return None, None

    col_emp, col_g, col_p, col_mail = st.columns([1.5, 1.5, 2, 2])

    empresas = sorted({g.get("empresa") or "—" for g in grupos})
    empresa_sel = col_emp.selectbox("Empresa", options=empresas, key=f"{prefix}_empresa")

    grupos_filtrados = [g for g in grupos if (g.get("empresa") or "—") == empresa_sel]
    if not grupos_filtrados:
        st.info("No hay grupos para esta empresa.")
        return None, None

    grupo_sel = col_g.selectbox(
        "Grupo", options=grupos_filtrados,
        format_func=lambda g: g["nombre"], key=f"{prefix}_grupo",
    )

    participantes = sorted(
        queries.listar_participantes(grupo_sel["id"]),
        key=lambda p: (p.get("pers_apellidos") or p.get("nombre") or "").lower(),
    )
    if not participantes:
        st.info("Este grupo no tiene participantes.")
        return None, None

    part_sel = col_p.selectbox(
        "Participante", options=participantes,
        format_func=lambda p: p["nombre"], key=f"{prefix}_part",
    )
    col_mail.text_input("Correo", value=part_sel.get("email") or "—",
                        disabled=True, key=f"{prefix}_mail_{part_sel['id']}")
    return grupo_sel, part_sel


def _tab_ingreso_auto():
    import pandas as pd
    grupo_sel, part_sel = _cascada_ingresos("iauto")
    if not grupo_sel or not part_sel:
        return

    if part_sel.get("autoevaluacion_completada"):
        st.warning("⚠️ Este participante ya tiene autoevaluación completada. Si guardas, se reemplazarán las respuestas existentes.")

    plantilla_id = grupo_sel.get("plantilla_id")
    if not plantilla_id:
        st.error("El grupo no tiene plantilla asignada.")
        return

    competencias = sorted(
        queries.listar_competencias_por_plantilla(plantilla_id),
        key=lambda c: c.get("orden", 0),
    )
    if not competencias:
        st.info("La plantilla no tiene competencias.")
        return

    st.divider()

    # Precargar respuestas existentes si las hay
    resp_existentes = queries.obtener_respuestas_participante(part_sel["id"]) or []
    resp_map = {r["competencia_id"]: r["puntaje"] for r in resp_existentes if r.get("es_autoevaluacion")}

    df = pd.DataFrame({
        "#": list(range(1, len(competencias) + 1)),
        "Competencia": [c["texto_auto"] for c in competencias],
        "Nota": [resp_map.get(c["id"], 3) for c in competencias],
    })

    edited = st.data_editor(
        df,
        column_config={
            "#": st.column_config.NumberColumn("#", disabled=True, width="small"),
            "Competencia": st.column_config.TextColumn("Competencia", disabled=True),
            "Nota": st.column_config.NumberColumn(
                "Nota", min_value=1, max_value=5, step=1, required=True, width="small",
            ),
        },
        hide_index=True,
        use_container_width=True,
        key=f"auto_editor_{part_sel['id']}",
    )

    if st.button("💾 Guardar Autoevaluación", type="primary", use_container_width=True, key="btn_guardar_auto"):
        try:
            scores = {competencias[i]["id"]: int(row["Nota"]) for i, row in edited.iterrows()}
            if part_sel.get("autoevaluacion_completada"):
                queries.eliminar_respuestas_auto(part_sel["id"])
            queries.guardar_respuestas_auto(part_sel["id"], scores)
            queries.actualizar_participante(part_sel["id"], {"autoevaluacion_completada": True})
            st.success(f"✅ Autoevaluación de {part_sel['nombre']} guardada correctamente.")
            st.rerun()
        except Exception as e:
            st.error(f"Error al guardar: {e}")


def _tab_ingreso_feedback():
    import pandas as pd
    grupo_sel, part_sel = _cascada_ingresos("ifb")
    if not grupo_sel or not part_sel:
        return

    plantilla_id = grupo_sel.get("plantilla_id")
    if not plantilla_id:
        st.error("El grupo no tiene plantilla asignada.")
        return

    competencias = queries.listar_competencias_por_plantilla(plantilla_id)
    if not competencias:
        st.info("La plantilla no tiene competencias.")
        return

    st.divider()

    # Mostrar evaluadores ya registrados para este participante
    evs = queries.listar_evaluadores(part_sel["id"])
    if not evs:
        st.info("Este participante no tiene evaluadores registrados.")
        return

    persona_ev_sel = st.selectbox(
        "Evaluador",
        options=[None] + evs,
        format_func=lambda e: (
            f"{e.get('nombre','—')} — {e.get('email','—')}"
            if e else "— Selecciona evaluador —"
        ),
        key="ifb_persona_sel",
    )

    if not persona_ev_sel:
        return

    ev_rut    = persona_ev_sel.get("pers_rut") or ""
    ev_nombre = persona_ev_sel.get("nombre", "")
    ev_email  = persona_ev_sel.get("email", "")

    st.divider()

    df = pd.DataFrame({
        "#": list(range(1, len(competencias) + 1)),
        "Competencia": [c["texto_feedback"] for c in competencias],
        "Nota": [3] * len(competencias),
    })

    edited = st.data_editor(
        df,
        column_config={
            "#": st.column_config.NumberColumn("#", disabled=True, width="small"),
            "Competencia": st.column_config.TextColumn("Competencia", disabled=True),
            "Nota": st.column_config.NumberColumn(
                "Nota", min_value=1, max_value=5, step=1, required=True, width="small",
            ),
        },
        hide_index=True,
        use_container_width=True,
        key=f"fb_editor_{part_sel['id']}",
    )

    if st.button("💾 Guardar Feedback", type="primary", use_container_width=True, key="btn_guardar_fb"):
        if not ev_nombre.strip() or not ev_email.strip():
            st.warning("Nombres, apellidos y correo del evaluador son obligatorios.")
        else:
            try:
                scores = {competencias[i]["id"]: int(row["Nota"]) for i, row in edited.iterrows()}
                evs_existentes = queries.listar_evaluadores(part_sel["id"])
                ev_existente = next(
                    (e for e in evs_existentes if e["email"].strip().lower() == ev_email.strip().lower()),
                    None,
                )
                if ev_existente:
                    ev_id = ev_existente["id"]
                    if ev_existente.get("completado"):
                        queries.eliminar_respuestas_feedback(part_sel["id"], ev_id)
                else:
                    pers_rut_fb = ev_rut.strip() if ev_rut.strip() else None
                    if not pers_rut_fb:
                        persona_fb = queries.buscar_persona_por_correo(ev_email.strip())
                        pers_rut_fb = persona_fb["pers_rut"] if persona_fb else None
                    ev_nuevo = queries.crear_evaluador(part_sel["id"], ev_nombre.strip(), ev_email.strip(), pers_rut_fb)
                    ev_id = ev_nuevo["id"]
                queries.guardar_respuestas_feedback(part_sel["id"], ev_id, scores)
                queries.actualizar_evaluador(ev_id, {"completado": True})
                st.success(f"✅ Feedback de {ev_nombre.strip()} para {part_sel['nombre']} guardado.")
                st.rerun()
            except Exception as e:
                st.error(f"Error al guardar: {e}")


def _tab_ingreso_evaluadores():
    import pandas as pd
    grupo_sel, part_sel = _cascada_ingresos("iev")
    if not grupo_sel or not part_sel:
        return

    st.divider()

    evs = queries.listar_evaluadores(part_sel["id"])
    enrolled_emails = {e["email"].strip().lower() for e in evs}

    # ---- Tabla editable de evaluadores actuales ----
    rows = []
    for i, ev in enumerate(evs, 1):
        rows.append({
            "✓": False,
            "#": i,
            "Nombre": ev.get("nombre") or "—",
            "Email": ev.get("email") or "—",
            "Contestó": "✅ Sí" if ev.get("completado") else "⏳ No",
            "_id": ev["id"],
        })

    st.markdown(f"**Evaluadores de {part_sel['nombre']}** ({len(evs)})")

    if rows:
        df_ev = pd.DataFrame(rows)
        edited_ev = st.data_editor(
            df_ev.drop(columns=["_id"]),
            column_config={
                "✓": st.column_config.CheckboxColumn("✓", default=False, width="small"),
                "#": st.column_config.NumberColumn("#", disabled=True, width="small"),
                "Nombre": st.column_config.TextColumn("Nombre"),
                "Email": st.column_config.TextColumn("Email"),
                "Contestó": st.column_config.TextColumn("Contestó", disabled=True, width="small"),
            },
            disabled=["#", "Contestó"],
            hide_index=True,
            use_container_width=True,
            key=f"ev_editor_{part_sel['id']}",
        )

        sel_idx = edited_ev.index[edited_ev["✓"] == True].tolist()
        col_grabar, col_eliminar, _ = st.columns([1, 1, 4])

        with col_grabar:
            if st.button("Grabar", use_container_width=True, type="primary", key=f"ev_grabar_{part_sel['id']}"):
                cambios = 0
                for i, row in edited_ev.iterrows():
                    orig = rows[i]
                    nuevo_nombre = row["Nombre"].strip()
                    nuevo_email  = row["Email"].strip()
                    if nuevo_nombre != orig["Nombre"] or nuevo_email != orig["Email"]:
                        queries.actualizar_evaluador(orig["_id"], {
                            "nombre": nuevo_nombre,
                            "email":  nuevo_email,
                        })
                        cambios += 1
                if cambios:
                    st.success(f"✅ {cambios} evaluador(es) actualizado(s).")
                    st.rerun()
                else:
                    st.info("Sin cambios para grabar.")

        with col_eliminar:
            if st.button(
                f"Eliminar ({len(sel_idx)})" if sel_idx else "Eliminar",
                use_container_width=True,
                disabled=(len(sel_idx) == 0),
                key=f"ev_eliminar_{part_sel['id']}",
            ):
                for i in sel_idx:
                    queries.eliminar_evaluador(rows[i]["_id"])
                st.success(f"✅ {len(sel_idx)} evaluador(es) eliminado(s).")
                st.rerun()
    else:
        st.info("Sin evaluadores registrados.")

    # ---- Agregar desde Personas ----
    st.divider()
    st.markdown("**Agregar evaluador desde Personas:**")
    rut_empresa = grupo_sel.get("rut_empresa")
    personas_disp = [
        p for p in queries.listar_personas_sist(rut_empresa=rut_empresa)
        if (p.get("pers_correo") or "").strip().lower() not in enrolled_emails
    ]
    if not personas_disp:
        personas_disp = [
            p for p in queries.listar_personas_sist()
            if (p.get("pers_correo") or "").strip().lower() not in enrolled_emails
        ]

    ca1, ca2 = st.columns([4, 1])
    persona_nueva = ca1.selectbox(
        "Persona",
        options=[None] + sorted(personas_disp, key=lambda p: (p.get("pers_apellidos") or "").lower()),
        format_func=lambda p: (
            f"{p['pers_apellidos']}, {p['pers_nombres']} — {p.get('pers_correo') or 'sin correo'}"
            if p else "— Selecciona persona —"
        ),
        key=f"ev_nueva_{part_sel['id']}",
    )
    with ca2:
        st.markdown("")
        if st.button("Agregar", use_container_width=True,
                     disabled=(persona_nueva is None),
                     key=f"ev_btn_agregar_{part_sel['id']}"):
            nombre = f"{persona_nueva['pers_nombres']} {persona_nueva['pers_apellidos']}".strip()
            queries.crear_evaluador(
                part_sel["id"], nombre,
                persona_nueva.get("pers_correo") or "",
                persona_nueva["pers_rut"],
            )
            st.success(f"✅ {nombre} agregado como evaluador.")
            st.rerun()


def pagina_ingresos_especiales():
    st.header("Ingresos Especiales")

    tab_ev, tab_auto, tab_fb = st.tabs([
        "Ingreso Evaluadores",
        "Ingreso Respuesta Autoevaluación",
        "Ingreso Respuestas Feedback",
    ])

    with tab_ev:
        _tab_ingreso_evaluadores()
    with tab_auto:
        _tab_ingreso_auto()
    with tab_fb:
        _tab_ingreso_feedback()


# ============================================================
# PÁGINA: MANTENEDORES
# ============================================================

def pagina_mantenedores():
    import pandas as pd
    st.header("Personas y Empresas")
    tab_personas, tab_empresas = st.tabs(["Personas", "Empresas"])

    # ----------------------------------------------------------
    # TAB PERSONAS (sist_personas)
    # ----------------------------------------------------------
    with tab_personas:
        empresas_otec = queries.listar_empresas_otec()
        empresa_map   = {e["rut_empresa"]: e["nombre_empresa"] for e in empresas_otec}
        emp_inv_map   = {e["nombre_empresa"]: e["rut_empresa"] for e in empresas_otec}
        emp_nombres   = ["— Sin empresa —"] + [e["nombre_empresa"] for e in empresas_otec]

        personas = queries.listar_personas_sist(solo_con_correo=False)

        if not personas:
            st.info("No hay personas registradas aún.")
        else:
            rows_p = []
            rut_list = []
            for p in personas:
                rows_p.append({
                    "✓": False,
                    "RUT": p.get("pers_rut", ""),
                    "Nombres": p.get("pers_nombres", "") or "",
                    "Apellidos": p.get("pers_apellidos", "") or "",
                    "Correo": p.get("pers_correo") or "",
                    "Empresa": empresa_map.get(p.get("rut_empresa") or "", "— Sin empresa —"),
                })
                rut_list.append(p.get("pers_rut", ""))

            df_p = pd.DataFrame(rows_p)
            edited_p = st.data_editor(
                df_p,
                use_container_width=True,
                hide_index=True,
                key="editor_personas",
                column_config={
                    "✓": st.column_config.CheckboxColumn("✓", default=False, width="small"),
                    "RUT": st.column_config.TextColumn("RUT", disabled=True),
                    "Nombres": st.column_config.TextColumn("Nombres"),
                    "Apellidos": st.column_config.TextColumn("Apellidos"),
                    "Correo": st.column_config.TextColumn("Correo"),
                    "Empresa": st.column_config.SelectboxColumn("Empresa", options=emp_nombres),
                },
            )
            st.caption(f"{len(personas)} persona(s) registrada(s).")

            col_g, col_e = st.columns(2)
            if col_g.button("Guardar cambios", type="primary", key="btn_save_personas"):
                errores = 0
                for i, row in edited_p.iterrows():
                    rut = rut_list[i]
                    emp_nombre = row["Empresa"]
                    rut_emp = emp_inv_map.get(emp_nombre) if emp_nombre != "— Sin empresa —" else None
                    try:
                        queries.actualizar_persona_sist(rut, {
                            "pers_nombres": row["Nombres"],
                            "pers_apellidos": row["Apellidos"],
                            "pers_correo": row["Correo"],
                            "rut_empresa": rut_emp,
                        })
                    except Exception:
                        errores += 1
                if errores:
                    st.error(f"Error al guardar {errores} registro(s).")
                else:
                    st.success("Cambios guardados correctamente.")
                    st.rerun()

            if col_e.button("Eliminar seleccionadas", key="btn_del_personas"):
                seleccionados = [rut_list[i] for i, row in edited_p.iterrows() if row["✓"]]
                if not seleccionados:
                    st.warning("Marca al menos una persona para eliminar.")
                else:
                    eliminados, con_error = [], []
                    for rut in seleccionados:
                        try:
                            queries.eliminar_persona_sist(rut)
                            eliminados.append(rut)
                        except Exception:
                            con_error.append(rut)
                    if eliminados:
                        st.success(f"{len(eliminados)} persona(s) eliminada(s).")
                    if con_error:
                        st.error(
                            f"{len(con_error)} persona(s) no se pudieron eliminar porque tienen "
                            "registros vinculados (participantes, evaluaciones, etc.): "
                            + ", ".join(con_error)
                        )
                    st.rerun()

        st.divider()

        with st.expander("➕ Crear nueva persona"):
            with st.form("form_persona"):
                col1, col2 = st.columns(2)
                with col1:
                    p_rut = st.text_input("RUT (sin puntos, con guión)", placeholder="12345678-9")
                    p_nombres = st.text_input("Nombres")
                    p_apellidos = st.text_input("Apellidos")
                with col2:
                    p_correo = st.text_input("Correo electrónico")
                    empresa_opts = [None] + empresas_otec
                    p_empresa_sel = st.selectbox(
                        "Empresa (opcional)",
                        options=empresa_opts,
                        format_func=lambda e: e["nombre_empresa"] if e else "— Sin empresa —",
                    )
                submitted = st.form_submit_button("Guardar persona", use_container_width=True, type="primary")
                if submitted:
                    if not p_rut.strip() or not p_nombres.strip() or not p_apellidos.strip() or not p_correo.strip():
                        st.warning("Completa RUT, Nombres, Apellidos y Correo.")
                    else:
                        rut_emp = p_empresa_sel["rut_empresa"] if p_empresa_sel else None
                        queries.crear_persona_sist(
                            p_rut.strip(), p_nombres.strip(), p_apellidos.strip(),
                            p_correo.strip(), rut_emp,
                        )
                        st.success(f"Persona {p_nombres.strip()} {p_apellidos.strip()} guardada.")
                        st.rerun()

    # ----------------------------------------------------------
    # TAB EMPRESAS
    # ----------------------------------------------------------
    with tab_empresas:
        empresas = queries.listar_empresas_otec()
        if not empresas:
            st.info("No hay empresas registradas.")
        else:
            rows_e = [{"RUT Empresa": e["rut_empresa"], "Nombre Empresa": e["nombre_empresa"]} for e in empresas]
            rut_e_list = [e["rut_empresa"] for e in empresas]
            df_e = pd.DataFrame(rows_e)
            edited_e = st.data_editor(
                df_e,
                use_container_width=True,
                hide_index=True,
                key="editor_empresas",
                column_config={
                    "RUT Empresa": st.column_config.TextColumn("RUT Empresa", disabled=True),
                    "Nombre Empresa": st.column_config.TextColumn("Nombre Empresa"),
                },
            )
            st.caption(f"{len(empresas)} empresa(s) registrada(s).")

            if st.button("Guardar cambios empresas", type="primary", key="btn_save_empresas"):
                errores = 0
                for i, row in edited_e.iterrows():
                    try:
                        queries.actualizar_empresa_otec(rut_e_list[i], {"nombre_empresa": row["Nombre Empresa"]})
                    except Exception:
                        errores += 1
                if errores:
                    st.error(f"Error al guardar {errores} registro(s).")
                else:
                    st.success("Cambios guardados correctamente.")
                    st.rerun()


# ============================================================
# ROUTER
# ============================================================

if menu == "Inicio":
    pagina_inicio()
elif menu == "Personas y Empresas":
    pagina_mantenedores()
elif menu == "Ingreso Encuestas":
    pagina_plantillas()
elif menu == "Ingreso de Grupos":
    pagina_grupos()
elif menu == "Ingresos Especiales":
    pagina_ingresos_especiales()
elif menu == "Seguimiento Autoevaluaciones":
    pagina_seguimiento_auto()
elif menu == "Seguimiento Feedback":
    pagina_seguimiento_feedback()
elif menu == "Informe Final 360":
    pagina_informe_final()
